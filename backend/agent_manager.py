# backend/agent_manager.py
import os
from typing import (
    Dict,
    Any,
    List,
    Tuple,
    Optional,
    Union,
    TypedDict,
    Set,
    AsyncGenerator,
)
from itertools import combinations
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_classic.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.output_parsers import StrOutputParser, PydanticOutputParser, JsonOutputParser
from langchain_core.messages import HumanMessage, AIMessage, BaseMessage, SystemMessage
from langchain_core.language_models.chat_models import (
    BaseChatModel,
)  # Import BaseChatModel
from langchain_openai import ChatOpenAI  # Import ChatOpenAI
from langchain_core.documents import Document
from datetime import datetime
import logging
from cachetools import TTLCache
import json
from asyncio import Lock, Event
from tenacity import retry, stop_after_attempt, wait_exponential
from langgraph.graph import StateGraph, END
from langchain_classic.chains.summarize import load_summarize_chain
from bs4 import BeautifulSoup  # Added import
from contextlib import asynccontextmanager
import asyncio
import re

# Removed SQLiteCache import
from langchain_classic.output_parsers import OutputFixingParser
from pydantic import BaseModel, Field, ValidationError
from datetime import timezone, timedelta
from api_key_manager import ApiKeyManager
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from database import (
    db_instance, User, Project, Chapter, CodexItem, CharacterVoiceProfile,
    CharacterRelationship, Event, Location, EventConnection, LocationConnection,
    KnowledgeBaseItem # Added
)
from vector_store import VectorStore
from graph_manager import GraphManager  # Added import
from models import (
    ChapterValidation,
    CodexItemBase as ModelCodexItem,
    WorldbuildingSubtype,
    CodexItemType,
    CodexExtractionTypes,
    RelationshipAnalysis,
    EventDescription,
    RelationshipAnalysisList,
    EventAnalysis,
    LocationConnection as LocationConnectionModel,
    LocationConnectionAnalysis,
    LocationAnalysis,
    LocationAnalysisList,
    EventConnectionBase,
    EventConnectionAnalysis,
    CodexExtraction,
    ProactiveSuggestionsResponse,
    ProactiveAssistRequest,
)
import io

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# Load environment variables
load_dotenv()

# --- Constants ---
PROCESS_TYPES = {
    "RELATIONSHIPS": "relationships",
    "LOCATIONS": "locations",
    "EVENTS": "events",
}
OPENROUTER_API_BASE = "https://openrouter.ai/api/v1"
SITE_URL = "https://github.com/LotusSerene/scrollwise-ai"
SITE_NAME = "ScrollWise AI"


# --- LangGraph State ---


class ChapterGenerationState(TypedDict):
    # Inputs
    user_id: str
    project_id: str
    chapter_number: int
    plot: str  # Represents the FULL plot for context retrieval
    writing_style: str
    instructions: Dict[
        str, Any
    ]  # Original instructions dict, may contain segment/total etc.

    # --- Added for Segmentation ---
    full_plot: Optional[str]  # Explicitly store the full plot
    plot_segment: Optional[str]  # Specific plot segment for this chapter (if any)
    total_chapters: int  # Total number of chapters being generated in this batch

    # --- For Batch Continuity ---
    # instructions["previous_chapters"] should contain previously generated chapters
    # from the same batch, in order, as a list of dicts with chapter_number, title, content

    # Dynamic values
    llm: BaseChatModel  # Use BaseChatModel type hint
    check_llm: BaseChatModel  # Use BaseChatModel type hint
    vector_store: VectorStore
    summarize_chain: Any  # Type hint could be improved

    # Intermediate results
    context: Optional[str] = None
    initial_chapter_content: Optional[str] = None
    extended_chapter_content: Optional[str] = None  # Store extension result separately
    current_word_count: int = 0
    target_word_count: int = 0
    chapter_title: Optional[str] = None
    new_codex_items: Optional[List[Dict[str, Any]]] = None
    validity_check: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    last_llm_response: Optional[BaseMessage] = (
        None  # Added for conversational continuation
    )
    project_description: Optional[str] = None  # Added for project description
    current_act_stage_info: Optional[str] = (
        None  # Added for current act/stage/substage info
    )

    # Final output (assembled at the end)
    final_chapter_content: Optional[str] = None
    final_output: Optional[Dict[str, Any]] = None


# --- Agent Manager ---


class AgentManager:
    _llm_cache = TTLCache(maxsize=100, ttl=3600)  # Class-level cache for LLM instances
    _graph_cache = {}  # Cache for compiled graphs per project

    def __init__(self, user_id: str, project_id: str, api_key_manager: ApiKeyManager):
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.DEBUG)  # Set higher level in production
        self.user_id = user_id
        self.project_id = project_id
        self.api_key_manager = api_key_manager
        self.model_settings = None
        self.api_key = None  # Gemini/Default key
        self.openrouter_api_key = None  # Add field for OpenRouter key
        self.anthropic_api_key = None  # Add field for Anthropic key
        self.openai_api_key = None  # Add field for direct OpenAI key
        self.MAX_INPUT_TOKENS = None
        self.vector_store = None
        self.graph_manager = GraphManager()  # Initialize GraphManager
        self.summarize_chain = None
        self.agents = {}  # Keep for potential future non-graph agents
        self._lock = Lock()  # Lock for managing shared resources like caches
        self.chapter_generation_graph = None  # Compiled LangGraph
        self.last_accessed = datetime.now(timezone.utc)  # Track last access time

    @classmethod
    async def create(
        cls, user_id: str, project_id: str, api_key_manager: ApiKeyManager
    ) -> "AgentManager":
        """Factory method to create and initialize AgentManager asynchronously."""
        key = (user_id, project_id)
        # Removed the 'if key in agent_managers:' check as the global dict is gone.

        instance = cls(user_id, project_id, api_key_manager)
        await instance.initialize()
        # Removed adding instance to global dict
        return instance

    async def initialize(self):
        """Initializes resources like API keys, LLMs, VectorStore, and the graph."""
        self.logger.info(
            f"Initializing AgentManager for User: {self.user_id[:8]}, Project: {self.project_id[:8]}"
        )
        try:
            # Fetch both API keys
            self.api_key = await self._get_api_key()  # Gemini key
            self.openrouter_api_key = await self.api_key_manager.get_openrouter_api_key(
                self.user_id
            )  # Fetch OpenRouter key
            self.anthropic_api_key = await self.api_key_manager.get_anthropic_api_key(
                self.user_id
            )  # Fetch Anthropic key
            self.openai_api_key = await self.api_key_manager.get_openai_api_key(
                self.user_id
            )  # Fetch OpenAI key

            if self.openrouter_api_key:
                self.logger.info(
                    f"OpenRouter API key loaded for user {self.user_id[:8]}."
                )
            else:
                self.logger.info(
                    f"No OpenRouter API key found for user {self.user_id[:8]}. OpenRouter models disabled."
                )
            if self.anthropic_api_key:
                self.logger.info(
                    f"Anthropic API key loaded for user {self.user_id[:8]}."
                )
            else:
                self.logger.info(
                    f"No Anthropic API key found for user {self.user_id[:8]}. Anthropic models disabled."
                )
            if self.openai_api_key:
                self.logger.info(f"OpenAI API key loaded for user {self.user_id[:8]}.")
            else:
                self.logger.info(
                    f"No OpenAI API key found for user {self.user_id[:8]}. Direct OpenAI models disabled."
                )

            self.model_settings = await self._get_model_settings()  # Changed to async

            # Determine token limits based on model type (simplified)
            # Need to refine this based on actual selected model later
            main_llm_name = self.model_settings.get("mainLLM", "")
            is_pro_model = "pro" in main_llm_name  # Very basic check
            if "gemini" in main_llm_name:
                self.MAX_INPUT_TOKENS = (
                    2_097_152 if is_pro_model else 1_048_576
                )  # Gemini 1.5 limits
            else:  # Add logic for OpenRouter model limits if needed, maybe default
                self.MAX_INPUT_TOKENS = 180000  # Increased from 32768

            self.setup_caching()  # Call caching setup

            # Initialize LLMs (using shared cache via _get_llm)
            # This now handles both Gemini and OpenRouter based on model name prefix
            self.llm = await self._get_llm(self.model_settings["mainLLM"])
            self.check_llm = await self._get_llm(self.model_settings["checkLLM"])

            # Initialize Vector Store (using Gemini embeddings key)
            self.vector_store = VectorStore(
                self.user_id,
                self.project_id,
                self.api_key,  # Use Gemini key for embeddings
                self.model_settings["embeddingsModel"],
            )
            
            # Check for migration OR empty collection (indicating failed previous migration or new project)
            # Safe to run even for empty projects as it just checks DB and adds nothing if DB is empty
            needs_migration = getattr(self.vector_store, "needs_migration", False)
            doc_count = await self.vector_store.get_count()
            
            if needs_migration or doc_count == 0:
                reason = "schema migration" if needs_migration else "empty collection"
                self.logger.info(f"Project {self.project_id} requires vector store update ({reason}). Starting reindexing...")
                await self._perform_vector_migration()
                self.logger.info(f"Project {self.project_id} migration/reindexing completed.")
            
            self.vector_store.set_llm(self.llm)  # Pass main LLM if needed by VS

            # Initialize Summarize Chain (using the appropriate LLM instance)
            self.summarize_chain = load_summarize_chain(
                self.llm, chain_type="map_reduce"
            )

            # Build and compile the chapter generation graph
            self.chapter_generation_graph = self._build_chapter_generation_graph()

            self.logger.info(
                f"AgentManager Initialized for User: {self.user_id[:8]}, Project: {self.project_id[:8]}"
            )

        except Exception as e:
            self.logger.error(f"Failed to initialize AgentManager: {e}", exc_info=True)
            raise  # Re-raise exception to indicate initialization failure

    async def _perform_vector_migration(self):
        """
        Fetches all project data and re-indexes it into the vector store.
        Used when a schema migration is detected.
        """
        try:
            async with db_instance.Session() as session:
                # 1. Fetch Codex Items
                codex_items_result = await session.execute(
                    select(CodexItem).where(
                        CodexItem.user_id == self.user_id,
                        CodexItem.project_id == self.project_id
                    )
                )
                codex_items = codex_items_result.scalars().all()
                
                # 2. Fetch Chapters
                chapters_result = await session.execute(
                    select(Chapter).where(
                        Chapter.user_id == self.user_id,
                        Chapter.project_id == self.project_id
                    )
                )
                chapters = chapters_result.scalars().all()

                # 3. Fetch Knowledge Base Items
                kb_items_result = await session.execute(
                    select(KnowledgeBaseItem).where(
                        KnowledgeBaseItem.user_id == self.user_id,
                        KnowledgeBaseItem.project_id == self.project_id
                    )
                )
                kb_items = kb_items_result.scalars().all()

                # Recreate Collection
                self.vector_store.recreate_collection()

                texts = []
                metadatas = []
                ids = []

                # Process Codex Items
                for item in codex_items:
                    text_content = f"{item.name}: {item.description}"
                    if item.backstory:
                        text_content += f"\nBackstory: {item.backstory}"
                    
                    metadata = {
                        "id": item.id,
                        "type": item.type,
                        "name": item.name,
                        "project_id": self.project_id,
                        "user_id": self.user_id,
                        "subtype": item.subtype
                    }
                    
                    texts.append(text_content)
                    metadatas.append(metadata)
                    ids.append(item.id)

                # Process Chapters
                for chapter in chapters:
                    if not chapter.content:
                        continue
                        
                    text_content = f"Chapter {chapter.chapter_number}: {chapter.title}\n{chapter.content}"
                    
                    metadata = {
                        "id": chapter.id,
                        "type": "chapter",
                        "title": chapter.title,
                        "chapter_number": chapter.chapter_number,
                        "project_id": self.project_id,
                        "user_id": self.user_id
                    }
                    
                    texts.append(text_content)
                    metadatas.append(metadata)
                    ids.append(chapter.id)

                # Process KB Items
                for kb_item in kb_items:
                    if not kb_item.content:
                        continue

                    metadata = kb_item.item_metadata or {}
                    metadata["id"] = kb_item.id
                    metadata["type"] = kb_item.type
                    metadata["project_id"] = self.project_id
                    metadata["user_id"] = self.user_id
                    metadata["source"] = kb_item.source
                    
                    texts.append(kb_item.content)
                    metadatas.append(metadata)
                    ids.append(kb_item.id)

                # Add to VectorStore
                if texts:
                    await self.vector_store.add_texts(texts=texts, metadatas=metadatas, ids=ids)
                    self.logger.info(f"Migrated {len(texts)} items for project {self.project_id}.")
                else:
                    self.logger.info(f"No items to migrate for project {self.project_id}.")

        except Exception as e:
            self.logger.error(f"Error during vector migration: {e}")
            raise

    async def close(self):
        """Cleans up resources like vector store connections."""
        self.logger.info(
            f"Closing AgentManager for User: {self.user_id[:8]}, Project: {self.project_id[:8]}"
        )
        # Clear LLM cache entries specific to this manager if needed (though TTLCache handles expiry)
        # self._llm_cache.clear() # Or selectively remove keys

        if self.vector_store:
            try:
                # Process any pending batched items before closing (if batching implemented)
                # if hasattr(self, '_embedding_batch') and self._embedding_batch:
                #    self.logger.info(f"Processing pending items before closing")
                #    await self._process_embedding_batch() # Assuming this method exists

                if hasattr(self.vector_store, "close") and callable(
                    self.vector_store.close
                ):
                    self.vector_store.close()  # Qdrant client might need closing
                    self.logger.debug("Vector store closed.")
            except Exception as e:
                self.logger.error(f"Error closing vector store: {e}", exc_info=True)

        # Clear graph cache if specific to this instance (not typically needed if stateless)
        # key = (self.user_id, self.project_id)
        # AgentManager._graph_cache.pop(key, None)

        # Removed removal from global tracking
        self.logger.info(
            f"AgentManager closed for User: {self.user_id[:8]}, Project: {self.project_id[:8]}"
        )

    async def _get_or_create_gemini_cache(
        self,
        cache_name: str,
        system_instruction: str,
        contents: List[str] = None,
        ttl_minutes: int = 60,
    ) -> str:
        """
        Creates or retrieves a Gemini context cache using google-genai SDK.
        Returns the cache name (resource name).
        """
        try:
            from google import genai
            from google.genai import types
            
            # Use the API key stored in the instance
            client = genai.Client(api_key=self.api_key)

            # List caches to find one with the matching display_name
            # Note: The new SDK listing might be synchronous or async depending on client configuration.
            # Assuming standard sync usage for now as per common migration patterns, or wrapping if needed.
            
            existing_cache = None
            # Using the list method which returns an iterable
            for c in client.caches.list():
                if c.display_name == cache_name:
                    existing_cache = c
                    break

            if existing_cache:
                self.logger.info(f"Found existing Gemini cache: {existing_cache.name}")
                # Update TTL
                client.caches.update(
                    name=existing_cache.name,
                    config=types.UpdateCachedContentConfig(
                        ttl=f"{ttl_minutes * 60}s" # TTL in seconds as string with 's' suffix
                    )
                )
                return existing_cache.name

            self.logger.info(f"Creating new Gemini cache: {cache_name}")
            
            # Create the cache
            cached_content = client.caches.create(
                model="models/gemini-flash-latest",
                config=types.CreateCachedContentConfig(
                    display_name=cache_name,
                    system_instruction=system_instruction,
                    contents=[types.Content(parts=[types.Part(text=c)]) for c in contents] if contents else [],
                    ttl=f"{ttl_minutes * 60}s",
                )
            )
            return cached_content.name

        except Exception as e:
            self.logger.error(f"Error creating Gemini cache: {e}")
            return None

    def _apply_anthropic_caching(self, messages: List[BaseMessage]) -> List[BaseMessage]:
        """
        Applies Anthropic's cache_control to the system message or the last large user message.
        """
        # If system message exists, cache it.
        # LangChain's ChatAnthropic handles system messages by looking for SystemMessage type.
        # We need to add additional_kwargs={"cache_control": {"type": "ephemeral"}} to it.
        
        updated_messages = []
        cache_applied = False
        
        # Check for SystemMessage first
        for msg in messages:
            if msg.type == "system" and not cache_applied:
                # Create a copy with cache_control
                new_kwargs = msg.additional_kwargs.copy()
                new_kwargs["cache_control"] = {"type": "ephemeral"}
                # Reconstruct the message (SystemMessage is a Pydantic model)
                # We can't easily modify it in place if it's immutable-ish, but we can create a new one.
                # Actually, just setting the dict might work if it's mutable.
                msg.additional_kwargs["cache_control"] = {"type": "ephemeral"}
                cache_applied = True
            updated_messages.append(msg)
            
        # If no system message, or if we want to cache the last large user context (like project structure)
        # For now, just caching the system prompt is a good start.
        
        return updated_messages

    @retry(
        stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10)
    )
    async def _get_llm(
        self, model_name: str, cached_content_name: Optional[str] = None
    ) -> BaseChatModel:  # Return BaseChatModel
        """Gets or creates a LangChain ChatModel instance (Gemini or OpenRouter) with caching."""
        
        cache_key = f"{model_name}_{cached_content_name}" if cached_content_name else model_name
        
        async with self._lock:  # Protect access to the class-level cache
            if cache_key in self._llm_cache:
                self.logger.debug(f"LLM Cache HIT for model: {model_name} (Key: {cache_key})")
                return self._llm_cache[cache_key]

            self.logger.debug(
                f"LLM Cache MISS for model: {model_name} (Key: {cache_key}). Creating new instance."
            )

            try:
                llm_instance: BaseChatModel

                # Check for OpenRouter prefix (e.g., "openrouter/openai/gpt-4o")
                if model_name.startswith("openrouter/"):
                    if not self.openrouter_api_key:
                        self.logger.error(
                            f"Attempted to use OpenRouter model '{model_name}' but OpenRouter API key is not set."
                        )
                        raise ValueError(
                            "OpenRouter API key is required for this model but not configured."
                        )

                    # Extract the actual model identifier (e.g., "openai/gpt-4o")
                    openrouter_model_id = model_name.split("openrouter/", 1)[1]

                    self.logger.info(
                        f"Initializing ChatOpenAI for OpenRouter model: {openrouter_model_id}"
                    )
                    # Configure ChatOpenAI for OpenRouter
                    llm_instance = ChatOpenAI(
                        model=openrouter_model_id,
                        openai_api_key=self.openrouter_api_key,
                        openai_api_base=OPENROUTER_API_BASE,
                        temperature=float(self.model_settings.get("temperature", 0.7)),
                        # max_tokens: ChatOpenAI might infer or have defaults. Can be set if needed.
                        # max_retries=2, # Example option
                        default_headers={  # Add required OpenRouter headers
                            "HTTP-Referer": SITE_URL,
                            "X-Title": SITE_NAME,
                        },
                        # Streaming is typically enabled by default or controlled by invoke/stream method
                    )
                    self.logger.info(
                        f"ChatOpenAI instance created for OpenRouter model: {openrouter_model_id}"
                    )

                # Check for Anthropic prefix (e.g., "anthropic/claude-3-opus-20240229")
                elif model_name.startswith("anthropic/"):
                    if not self.anthropic_api_key:
                        self.logger.error(
                            f"Attempted to use Anthropic model '{model_name}' but Anthropic API key is not set."
                        )
                        raise ValueError(
                            "Anthropic API key is required for this model but not configured."
                        )

                    anthropic_model_id = model_name.split("anthropic/", 1)[1]
                    self.logger.info(
                        f"Initializing ChatAnthropic for model: {anthropic_model_id}"
                    )
                    from langchain_anthropic import ChatAnthropic

                    llm_instance = ChatAnthropic(
                        model=anthropic_model_id,
                        anthropic_api_key=self.anthropic_api_key,
                        temperature=float(self.model_settings.get("temperature", 0.7)),
                    )
                    self.logger.info(
                        f"ChatAnthropic instance created for model: {anthropic_model_id}"
                    )

                # Check for direct OpenAI prefix (e.g., "openai/gpt-4o")
                elif model_name.startswith("openai/"):
                    if not self.openai_api_key:
                        self.logger.error(
                            f"Attempted to use direct OpenAI model '{model_name}' but OpenAI API key is not set."
                        )
                        raise ValueError(
                            "OpenAI API key is required for this model but not configured."
                        )

                    openai_model_id = model_name.split("openai/", 1)[1]
                    self.logger.info(
                        f"Initializing ChatOpenAI for direct model: {openai_model_id}"
                    )
                    llm_instance = ChatOpenAI(
                        model=openai_model_id,
                        openai_api_key=self.openai_api_key,
                        temperature=float(self.model_settings.get("temperature", 0.7)),
                    )
                    self.logger.info(
                        f"ChatOpenAI instance created for direct model: {openai_model_id}"
                    )

                # Default to Gemini if no specific prefix matches
                else:
                    if not self.api_key:
                        self.logger.error(
                            f"Attempted to use Gemini model '{model_name}' but Gemini API key is not set."
                        )
                        raise ValueError(
                            "Gemini API key is required for this model but not configured."
                        )

                    self.logger.info(f"Initializing ChatGoogleGenerativeAI for model: {model_name}")
                    llm_instance = ChatGoogleGenerativeAI(
                        model=model_name,
                        google_api_key=self.api_key,
                        temperature=float(self.model_settings.get("temperature", 0.7)),
                        convert_system_message_to_human=True,
                        cached_content=cached_content_name,
                    )
                    self.logger.info(
                        f"ChatGoogleGenerativeAI instance created for model: {model_name}"
                    )

                # Cache the created instance
                self._llm_cache[cache_key] = llm_instance
                return llm_instance

            except ValueError as ve:  # Catch specific configuration errors
                self.logger.error(
                    f"Configuration error creating LLM instance for {model_name}: {ve}"
                )
                raise
            except Exception as e:
                self.logger.error(
                    f"Failed to create LLM instance for {model_name}: {e}",
                    exc_info=True,
                )
                raise  # Re-raise the exception after logging

    async def _get_api_key(self) -> str:
        """Gets the primary (Gemini) API key."""  # Updated docstring
        api_key = await self.api_key_manager.get_api_key(self.user_id)
        if not api_key:
            self.logger.error(
                f"Primary (Gemini) API key not found for user {self.user_id}"
            )
            raise ValueError(
                "Primary API key not set. Please set your Gemini API key in the settings."
            )
        return api_key

    async def _get_model_settings(self) -> dict:
        """Fetches model settings from the database asynchronously."""
        try:
            settings = await db_instance.get_model_settings(self.user_id)
            # Provide defaults carefully, especially for models
            # Defaulting to Gemini Flash if nothing is set
            defaults = {
                "mainLLM": "gemini-1.5-flash-latest",
                "checkLLM": "gemini-1.5-flash-latest",
                "embeddingsModel": "models/gemini-embedding-001",  # Keep Gemini embedding for now
                "titleGenerationLLM": "gemini-1.5-flash-latest",
                "extractionLLM": "gemini-1.5-flash-latest",
                "knowledgeBaseQueryLLM": "gemini-1.5-flash-latest",
                "temperature": 0.7,
            }
            # Ensure loaded settings overwrite defaults
            final_settings = {**defaults}  # Start with defaults
            if settings:  # Check if settings were actually loaded
                final_settings.update(
                    {k: v for k, v in settings.items() if v is not None}
                )  # Merge non-None loaded settings

            final_settings["temperature"] = float(
                final_settings["temperature"]
            )  # Ensure float
            self.logger.debug(
                f"Final model settings for user {self.user_id[:8]}: {final_settings}"
            )
            return final_settings
        except Exception as e:
            self.logger.error(
                f"Error getting model settings for user {self.user_id}: {e}",
                exc_info=True,
            )
            raise

    def setup_caching(self):
        """Disables local file-based LLM caching for cloud compatibility."""
        # set_llm_cache(None) # Explicitly disable caching if needed
        # Or simply do nothing, as SQLiteCache is no longer imported/used
        self.logger.info("Local file-based LLM caching disabled for cloud deployment.")
        # Note: Consider external caching (Redis/Memcached) for performance in production.

    def estimate_token_count(self, text: str) -> int:
        """Estimates token count using the main LLM."""
        if not text or not isinstance(text, str):
            return 0
        try:
            # Use the initialized main LLM instance
            return self.llm.get_num_tokens(text)
        except Exception as e:
            self.logger.warning(
                f"Could not estimate token count: {e}. Using word count approximation."
            )
            return len(text.split())  # Fallback approximation

    def _format_project_structure(
        self,
        structure_data: Any,
        level: int = 0,
    ) -> str:
        """
        Recursively formats the project structure (acts, stages, substages, chapters) as a readable outline.
        Handles both dictionary with 'project_structure' key or direct list of structure items.
        """
        if not structure_data:
            return ""

        # Handle both direct list or dict with 'project_structure' key
        items = []
        if isinstance(structure_data, dict) and "project_structure" in structure_data:
            items = structure_data.get("project_structure", [])
        elif isinstance(structure_data, list):
            items = structure_data

        if not isinstance(items, list):
            self.logger.warning(
                f"Project structure data is not in the expected format for project {self.project_id}"
            )
            return ""

        lines = []
        for item in items:
            lines.extend(self._format_single_structure_item(item, level))

        return "\n".join(lines)

    def _format_single_structure_item(self, item_data: Dict, level: int) -> List[str]:
        """Formats a single structure item (act, stage, substage, or chapter) from a dictionary representation."""
        lines = []
        indent = "    " * level  # 4 spaces per indent level

        item_id = item_data.get("id")
        item_type = item_data.get("type", "unknown")
        # Handle both name (backend) and title (frontend) fields
        item_name = item_data.get("name", item_data.get("title", "Unnamed Item"))
        item_description = item_data.get("description", "")

        # Format based on type
        if item_type == "chapter":
            # Chapter format
            chapter_number = item_data.get("chapter_number", "?")
            lines.append(
                f"{indent}- Chapter {chapter_number}: {item_name} (ID: {item_id})"
            )
        else:  # act, stage, substage, folder
            # Structure item format
            lines.append(f"{indent}+ {item_type.capitalize()}: {item_name}")
            if item_description:
                lines.append(f"{indent}  Description: {item_description}")

        # Process children (can be folders or chapters)
        children = item_data.get("children", [])
        if children and isinstance(children, list):
            for child in children:
                if isinstance(child, dict):
                    lines.extend(self._format_single_structure_item(child, level + 1))
                else:
                    self.logger.warning(f"Child item is not a dict, skipping: {child}")

        return lines

    def _find_structure_item_name(
        self, project_structure: Any, item_id: str
    ) -> Optional[str]:
        """
        Recursively finds the name of a structure item (act, stage, substage) by its ID.
        Returns a breadcrumb-like string if found (e.g., "Act I > Stage 1 > Substage A").
        """
        if not project_structure or not isinstance(project_structure, list):
            return None

        for item in project_structure:
            if isinstance(item, dict):  # Ensure item is a dictionary
                current_name = item.get("name", "Unnamed Item")
                if item.get("id") == item_id:
                    return current_name
                if "children" in item and item["children"]:
                    child_path = self._find_structure_item_name(
                        item["children"], item_id
                    )
                    if child_path:
                        return f"{current_name} - {child_path}"
            elif hasattr(item, "id") and hasattr(
                item, "name"
            ):  # Handle Pydantic models
                current_name = item.name if item.name else "Unnamed Item"
                if item.id == item_id:
                    return current_name
                if hasattr(item, "children") and item.children:
                    child_path = self._find_structure_item_name(item.children, item_id)
                    if child_path:
                        return f"{current_name} - {child_path}"
        return None

    def _find_structure_item_details(
        self, project_structure: Any, item_id: str
    ) -> List[Dict[str, Optional[str]]]:
        """
        Recursively finds the path of structure items (act, stage, substage, chapter) by ID,
        returning a list of dictionaries with 'name' and 'description'.
        Works with both old and new structure formats.
        """
        if not project_structure or not isinstance(project_structure, list):
            return []

        path = []

        def find_path_recursive(
            structure_list: List[Any], current_item_id: str
        ) -> bool:
            for item_data in structure_list:
                item = {}
                children_list = []

                if isinstance(item_data, dict):
                    # Handle both name (backend) and title (frontend) fields
                    item = {
                        "id": item_data.get("id"),
                        "name": item_data.get(
                            "name", item_data.get("title", "Unnamed")
                        ),
                        "description": item_data.get("description"),
                        "type": item_data.get("type", "unknown"),
                    }
                    children_list = item_data.get("children", [])
                elif hasattr(item_data, "id") and hasattr(
                    item_data, "name"
                ):  # Handle Pydantic models
                    item = {
                        "id": item_data.id,
                        "name": getattr(
                            item_data, "name", getattr(item_data, "title", "Unnamed")
                        ),
                        "description": getattr(item_data, "description", None),
                        "type": getattr(item_data, "type", "unknown"),
                    }
                    children_list = getattr(item_data, "children", [])

                if not item.get("id"):
                    continue

                path.append(
                    {
                        "name": item["name"],
                        "description": item["description"],
                        "type": item["type"],
                    }
                )

                if item["id"] == current_item_id:
                    return True

                if children_list:
                    if find_path_recursive(children_list, current_item_id):
                        return True

                path.pop()  # Backtrack
            return False

        find_path_recursive(project_structure, item_id)
        return path

    def _normalize_name(self, name: str) -> str:
        """Standardize name for comparison by removing extra spaces, punctuation and converting to lowercase."""
        if not name or not isinstance(name, str):
            return ""
        # Remove punctuation, convert to lowercase, and collapse whitespace.
        name = name.lower()
        name = re.sub(r"[^\w\s-]", "", name)  # Allow words, spaces, hyphens
        name = re.sub(r"\s+", " ", name).strip()
        return name

    # --- LangGraph Nodes ---

    async def _construct_context_node(
        self, state: ChapterGenerationState
    ) -> Dict[str, Any]:
        """Gathers context for chapter generation."""
        self.logger.debug(
            f"Node: Constructing Context for Chapter {state['chapter_number']}"
        )
        try:
            plot = state["plot"]  # This should be the full plot from initial call
            writing_style = state["writing_style"]
            vector_store = state["vector_store"]
            user_id = state["user_id"]
            project_id = state["project_id"]
            chapter_number = state[
                "chapter_number"
            ]  # Get chapter_number for fetching chapter details

            # Get any previous chapters from the current batch if available
            instructions = state["instructions"]
            previous_chapters_from_batch = instructions.get("previous_chapters", [])

            context_parts = []
            project_description_for_state: Optional[str] = None
            current_act_stage_info_for_state: Optional[str] = None
            full_project_structure_formatted: Optional[str] = None  # New variable

            # Fetch project details (description and structure)
            project_details = await db_instance.get_project(
                user_id=user_id, project_id=project_id
            )

            if project_details:
                project_description_for_state = project_details.get("description")
                project_structure_json = project_details.get("project_structure")
                # self.logger.debug(f"Fetched project_structure_json: {project_structure_json} (Type: {type(project_structure_json)})") # <<< REMOVED DEBUG LOG

                # Fetch all chapters to build a map for linking
                all_chapters_data = await db_instance.get_all_chapters(
                    user_id, project_id
                )
                chapters_by_structure_id_map = {
                    ch["structure_item_id"]: ch
                    for ch in all_chapters_data
                    if ch.get("structure_item_id")
                }

                if project_structure_json:
                    # Normalize the structure to always be a list of items for processing
                    structure_list = []
                    if isinstance(project_structure_json, dict):
                        structure_list = project_structure_json.get(
                            "project_structure", []
                        )
                    elif isinstance(project_structure_json, list):
                        structure_list = project_structure_json

                    # Format the ENTIRE project structure for the LLM context
                    # The helper function `_format_project_structure` can handle both old and new formats.
                    full_project_structure_formatted = self._format_project_structure(
                        project_structure_json
                    )
                    if full_project_structure_formatted:
                        self.logger.info(
                            "Formatted full project structure for LLM context."
                        )

                    # Get current chapter's linked structure item ID
                    current_chapter_db_info = await db_instance.get_chapter_by_number(
                        project_id=project_id,
                        user_id=user_id,
                        chapter_number=chapter_number,
                    )
                    chapter_structure_item_id = None
                    if current_chapter_db_info:
                        chapter_structure_item_id = current_chapter_db_info.get(
                            "structure_item_id"
                        )

                    if chapter_structure_item_id:
                        # Find the path to the current chapter in the structure
                        item_details_path = self._find_structure_item_details(
                            structure_list,  # Use the normalized list here
                            chapter_structure_item_id,
                        )
                        if item_details_path:
                            path_components = []
                            for item in item_details_path:
                                name = item.get("name", "Unnamed Item")
                                description = item.get("description")
                                item_type = item.get("type", "unknown")
                                component = f"{item_type.capitalize()}: {name}"
                                if description:
                                    component += f" ({description})"
                                path_components.append(component)
                            current_act_stage_info_for_state = (
                                "Current Structural Context for this Chapter: "
                                + " > ".join(path_components)
                            )
                            self.logger.info(
                                f"Added current structural context: {current_act_stage_info_for_state}"
                            )
            # Add plot and writing style to general context_parts
            context_parts.append(f"Plot: {plot}")
            context_parts.append(f"Writing Style: {writing_style}")

            # Add the current structural context to general context if available
            if current_act_stage_info_for_state:
                context_parts.append(current_act_stage_info_for_state)

            # Add the FULL project structure to context_parts if available
            if full_project_structure_formatted:
                context_parts.append(
                    f"\n**FULL PROJECT STRUCTURE (Acts, Stages, Substages, Linked Chapters):**\n{full_project_structure_formatted}"
                )

            # Fetch user subscription status
            is_pro_user_for_voice = False
            subscription_info = await db_instance.get_user_subscription_info(user_id)
            if subscription_info:
                is_pro_user_for_voice = (
                    subscription_info.get("plan") == "pro"
                    and subscription_info.get("status") == "active"
                )
                self.logger.info(
                    f"User {user_id} subscription for voice profile feature: Plan='{subscription_info.get("plan")}', Status='{subscription_info.get("status")}'. Voice Active: {is_pro_user_for_voice}"
                )
            else:
                self.logger.warning(
                    f"User {user_id} subscription info not found. Voice profiles disabled for this run."
                )

            # Check if we have previous chapters from the current batch
            if previous_chapters_from_batch:
                self.logger.info(
                    f"Found {len(previous_chapters_from_batch)} chapters from the current generation batch"
                )
                context_parts.append("\nPreviously Generated Chapters In This Batch:")

                # Always include full content of recently generated chapters in the same batch
                # These are the most important for maintaining narrative continuity
                for ch in previous_chapters_from_batch:
                    ch_num = ch.get("chapter_number", "?")
                    ch_title = ch.get("title", f"Chapter {ch_num}")
                    ch_content = ch.get("content", "")
                    context_parts.append(
                        f"\nCHAPTER {ch_num}: {ch_title}\n\n{ch_content}\n"
                    )
                    self.logger.info(
                        f"Added Chapter {ch_num} from current batch to context"
                    )

            # Fetch relevant previous chapters from database (for older chapters)
            try:
                previous_chapters_data = await db_instance.get_all_chapters(
                    user_id, project_id
                )
                previous_chapters_data.sort(key=lambda x: x.get("chapter_number", 0))

                # Implement smart rolling memory with progressive summarization
                if previous_chapters_data:
                    # Don't include the current chapter or any chapters we already included from the batch
                    batch_chapter_numbers = [
                        ch.get("chapter_number") for ch in previous_chapters_from_batch
                    ]
                    previous_chapters_data = [
                        ch
                        for ch in previous_chapters_data
                        if ch.get("chapter_number") != chapter_number
                        and ch.get("chapter_number") not in batch_chapter_numbers
                    ]

                    if previous_chapters_data:
                        context_parts.append("\nPrevious Chapter Information:")

                        # First, try to estimate total token count for all chapters
                        chapter_token_estimates = []
                        total_chapter_tokens = 0

                        for ch in previous_chapters_data:
                            ch_content = ch.get("content", "")
                            ch_tokens = self.estimate_token_count(ch_content)
                            chapter_token_estimates.append((ch, ch_tokens))
                            total_chapter_tokens += ch_tokens

                        # Maximum tokens to allow for chapters (reserving space for other context)
                        max_chapters_tokens = self.MAX_INPUT_TOKENS // 2
                        self.logger.info(
                            f"Previous chapters total: {total_chapter_tokens} tokens, limit: {max_chapters_tokens}"
                        )

                        # If all chapters fit, use individual summaries for long ones
                        if total_chapter_tokens <= max_chapters_tokens:
                            self.logger.info(
                                "Including all chapters with individual processing"
                            )
                            for ch, ch_tokens in chapter_token_estimates:
                                chap_num = ch.get("chapter_number", "N/A")
                                chap_title = ch.get("title", f"Chapter {chap_num}")
                                content = ch.get("content", "")

                                # Summarize individual chapters if they're long
                                if ch_tokens > 2500:
                                    docs_to_summarize = [Document(page_content=content)]
                                    summary_result = await state[
                                        "summarize_chain"
                                    ].ainvoke({"input_documents": docs_to_summarize})
                                    context_parts.append(
                                        f"- Ch {chap_num} ({chap_title}): {summary_result.get('output_text', 'Summary unavailable')}"
                                    )
                                else:
                                    context_parts.append(
                                        f"- Ch {chap_num} ({chap_title}): {content[:2100]}..."
                                    )
                        else:
                            # Progressive summarization approach when all chapters don't fit
                            self.logger.info(
                                "Using progressive summarization for chapters"
                            )

                            # Always include the most recent chapters (last 3) with individual treatment
                            recent_chapters = previous_chapters_data[-3:]
                            older_chapters = previous_chapters_data[:-3]

                            # Process recent chapters individually (last 3)
                            for ch in recent_chapters:
                                chap_num = ch.get("chapter_number", "N/A")
                                chap_title = ch.get("title", f"Chapter {chap_num}")
                                content = ch.get("content", "")

                                # Summarize individual recent chapters if they're long
                                if self.estimate_token_count(content) > 2500:
                                    docs_to_summarize = [Document(page_content=content)]
                                    summary_result = await state[
                                        "summarize_chain"
                                    ].ainvoke({"input_documents": docs_to_summarize})
                                    context_parts.append(
                                        f"- Ch {chap_num} ({chap_title}): {summary_result.get('output_text', 'Summary unavailable')}"
                                    )
                                else:
                                    context_parts.append(
                                        f"- Ch {chap_num} ({chap_title}): {content[:2100]}..."
                                    )

                            # If we have older chapters, summarize them in batches
                            if older_chapters:
                                # Group older chapters into batches of 10
                                batch_size = 10
                                chapter_batches = []

                                for i in range(0, len(older_chapters), batch_size):
                                    chapter_batches.append(
                                        older_chapters[i : i + batch_size]
                                    )

                                self.logger.info(
                                    f"Processing {len(chapter_batches)} batches of older chapters"
                                )

                                for batch_idx, batch in enumerate(chapter_batches):
                                    # Create a combined document for the batch
                                    batch_start = batch[0].get("chapter_number", "?")
                                    batch_end = batch[-1].get("chapter_number", "?")
                                    batch_content = "\n\n---\n\n".join(
                                        [
                                            f"Chapter {ch.get('chapter_number')}: {ch.get('title', '')}\n{ch.get('content', '')}"
                                            for ch in batch
                                        ]
                                    )

                                    # Summarize the batch
                                    docs_to_summarize = [
                                        Document(page_content=batch_content)
                                    ]
                                    batch_summary = await state[
                                        "summarize_chain"
                                    ].ainvoke({"input_documents": docs_to_summarize})

                                    context_parts.append(
                                        f"- Chapters {batch_start}-{batch_end} Summary: {batch_summary.get('output_text', 'Batch summary unavailable')}"
                                    )
            except Exception as e:
                self.logger.warning(f"Could not fetch/process previous chapters: {e}")


            relevant_docs = []
            try:
                # More specific query for relevance
                query_text = f"Details relevant to plot: {plot}"
                # Filter out chapters explicitly
                codex_filter = {
                    "type": {"$nin": ["chapter", "relationship", "character_backstory"]}
                }  # Example filter

                relevant_docs = await vector_store.similarity_search(
                    query_text=query_text,
                    k=20,  # Fetch more potentially relevant items
                    filter=codex_filter,
                )

                relevant_entity_names = []  # Collect names for Graph Context

                if relevant_docs:
                    context_parts.append(
                        "\nRelevant World Information (from Knowledge Base):"
                    )
                    items_by_type = {}
                    for doc in relevant_docs:
                        item_type = doc.metadata.get("type", "other")
                        if item_type not in items_by_type:
                            items_by_type[item_type] = []
                        # Limit context per item to avoid excessive length
                        name = doc.metadata.get(
                            "name", doc.metadata.get("title", "Unnamed")
                        )
                        relevant_entity_names.append(name)

                        content_preview = doc.page_content[:300]  # Limit length
                        items_by_type[item_type].append(f"{name}: {content_preview}...")

                    for (
                        item_type_key,
                        items_list,
                    ) in items_by_type.items():  # Renamed to avoid conflict
                        context_parts.append(f"\n{item_type_key.title()}:")
                        context_parts.extend([f"- {item}" for item in items_list])

                    # --- GRAPH ENHANCED CONTEXT ---
                    try:
                        self.logger.debug("Building Graph Context...")
                        async with db_instance.Session() as session:
                            # Fetch Codex Items
                            codex_items_result = await session.execute(
                                select(CodexItem).options(selectinload(CodexItem.voice_profile)).where(CodexItem.project_id == project_id)
                            )
                            codex_items_data = [item.to_dict() for item in codex_items_result.scalars().all()]

                            # Fetch Relationships
                            rels_result = await session.execute(
                                select(CharacterRelationship).where(CharacterRelationship.project_id == project_id)
                            )
                            rels_data = [rel.to_dict() for rel in rels_result.scalars().all()]

                            # Fetch Events
                            events_result = await session.execute(
                                select(Event).where(Event.project_id == project_id)
                            )
                            events_data = [evt.to_dict() for evt in events_result.scalars().all()]

                            # Fetch Locations
                            locs_result = await session.execute(
                                select(Location).where(Location.project_id == project_id)
                            )
                            locs_data = [loc.to_dict() for loc in locs_result.scalars().all()]

                            # Fetch Event Connections
                            evt_conns_result = await session.execute(
                                select(EventConnection).where(EventConnection.project_id == project_id)
                            )
                            evt_conns_data = [conn.to_dict() for conn in evt_conns_result.scalars().all()]

                            # Fetch Location Connections
                            loc_conns_result = await session.execute(
                                select(LocationConnection).where(LocationConnection.project_id == project_id)
                            )
                            loc_conns_data = [conn.to_dict() for conn in loc_conns_result.scalars().all()]

                        # Build Graph
                        self.graph_manager.build_graph(
                            codex_items=codex_items_data,
                            relationships=rels_data,
                            events=events_data,
                            locations=locs_data,
                            event_connections=evt_conns_data,
                            location_connections=loc_conns_data
                        )

                        # Get Related Context
                        graph_context = self.graph_manager.get_related_context(relevant_entity_names)
                        if graph_context:
                            context_parts.append("\n" + graph_context)
                            self.logger.debug("Added Graph Context.")
                        else:
                            self.logger.debug("No Graph Context found.")

                    except Exception as graph_e:
                        self.logger.warning(f"Failed to generate Graph Context: {graph_e}")
                    # -----------------------------

                # Conditionally add character voice profiles for Pro users
                if is_pro_user_for_voice:
                    character_voice_profiles_context_list = []  # Renamed variable
                    self.logger.debug(
                        f"Pro user {user_id}: Attempting to fetch voice profiles for relevant characters."
                    )

                    character_docs = [
                        doc
                        for doc in relevant_docs
                        if doc.metadata.get("type") == CodexItemType.CHARACTER.value
                    ]
                    if character_docs:
                        processed_character_names = set()
                        temp_profile_strings = []
                        for doc in character_docs:
                            character_name = doc.metadata.get(
                                "name", "Unknown Character"
                            )
                            # Ensure we process each character only once, even if multiple docs refer to them
                            if character_name in processed_character_names:
                                continue
                            processed_character_names.add(character_name)

                            codex_item_id = doc.metadata.get("db_item_id")

                            if codex_item_id:
                                voice_profile_dict = await db_instance.get_character_voice_profile_by_codex_id(
                                    codex_item_id=codex_item_id,
                                    user_id=user_id,
                                    project_id=project_id,
                                )
                                if voice_profile_dict:
                                    profile_details = [
                                        f"Character: {character_name}"
                                    ]  # Removed ID for brevity in prompt
                                    if voice_profile_dict.get("vocabulary"):
                                        profile_details.append(
                                            f"  - Vocabulary: {voice_profile_dict['vocabulary']}"
                                        )
                                    if voice_profile_dict.get("sentence_structure"):
                                        profile_details.append(
                                            f"  - Sentence Structure: {voice_profile_dict['sentence_structure']}"
                                        )
                                    if voice_profile_dict.get("speech_patterns_tics"):
                                        profile_details.append(
                                            f"  - Speech Patterns/Tics: {voice_profile_dict['speech_patterns_tics']}"
                                        )
                                    if voice_profile_dict.get("tone"):
                                        profile_details.append(
                                            f"  - Tone: {voice_profile_dict['tone']}"
                                        )
                                    if voice_profile_dict.get("habits_mannerisms"):
                                        profile_details.append(
                                            f"  - Habits/Mannerisms: {voice_profile_dict['habits_mannerisms']}"
                                        )

                                    if (
                                        len(profile_details) > 1
                                    ):  # Found some actual voice data
                                        temp_profile_strings.append(
                                            "\n".join(profile_details)
                                        )
                                        self.logger.debug(
                                            f"Formatted voice profile for character: {character_name}"
                                        )

                        if temp_profile_strings:
                            character_voice_profiles_context_list.append(
                                "\n"
                            )
                            character_voice_profiles_context_list.extend(
                                temp_profile_strings
                            )
                            context_parts.extend(character_voice_profiles_context_list)
                    else:
                        self.logger.debug(
                            "No character type documents found in relevant_docs for voice profile processing."
                        )
                else:
                    self.logger.info(
                        f"User {user_id} is not Pro or subscription not active. Skipping voice profile additions."
                    )

            except Exception as e:
                self.logger.warning(
                    f"Could not fetch relevant documents from vector store: {e}"
                )

            final_context = "\n".join(context_parts)

            # Truncate context if it exceeds limits (leaving room for prompt template)
            max_context_tokens = (
                self.MAX_INPUT_TOKENS
                - 1500  # Reserve 1.5k for prompt itself and other state fields
            )
            context_tokens = self.estimate_token_count(final_context)
            if context_tokens > max_context_tokens:
                self.logger.warning(
                    f"Context ({context_tokens} tokens) exceeds limit ({max_context_tokens}). Truncating."
                )
                ratio = max_context_tokens / context_tokens
                final_context = final_context[: int(len(final_context) * ratio)]
                self.logger.warning(
                    f"Context truncated to approx {self.estimate_token_count(final_context)} tokens."
                )

            return {
                "context": final_context,
                "project_description": project_description_for_state,
                "current_act_stage_info": current_act_stage_info_for_state,
            }

        except Exception as e:
            self.logger.error(f"Error in _construct_context_node: {e}", exc_info=True)
            return {"error": f"Failed to construct context: {e}"}

    def _create_chapter_prompt(
        self, state: ChapterGenerationState  # Accept the full state
    ) -> ChatPromptTemplate:
        """Helper to create the chapter generation prompt, incorporating plot segmentation,
        project description, current act/stage information, and full project structure conditionally.
        """

        # Extract relevant fields from state
        instructions = state["instructions"]
        context = state[
            "context"
        ]  # This is the general context (codex, previous chaps, full structure etc)
        chapter_number = state["chapter_number"]
        total_chapters = state["total_chapters"]
        full_plot = state["full_plot"]
        plot_segment = state.get("plot_segment")
        writing_style = state["writing_style"]
        style_guide = instructions.get("styleGuide", "")
        additional_instructions_text = instructions.get(
            "additionalInstructions", ""
        )  # Renamed to avoid conflict
        word_count_target = instructions.get("wordCount", 0)
        project_description = state.get("project_description")
        # current_act_stage_info is now part of the main {{context}}
        # The full project structure is also part of {{context}}

        # Determine the focus plot for the prompt
        focus_plot_description = plot_segment if plot_segment else full_plot
        plot_header = (
            "CURRENT CHAPTER PLOT SEGMENT (Main Focus for this Chapter):"
            if plot_segment
            else "PLOT/SETTING (Focus for this Chapter):"
        )

        # Build optional prompt sections
        optional_prompt_parts = []
        if project_description:
            optional_prompt_parts.append(
                f"**PROJECT DESCRIPTION (Overall Theme/Goal):**\\n{{project_description}}"
            )
        # current_act_stage_info and full_project_structure are now part of {{context}}

        optional_prompt_section = (
            "\\n\\n" + "\\n\\n".join(optional_prompt_parts)
            if optional_prompt_parts
            else ""
        )

        system_template = f"""You are a skilled author writing chapter {{chapter_number}} of {{total_chapters}} for a novel. Your goal is to write in a style that is engaging, natural, and human-like. Adhere STRICTLY to all requirements.
    
        **HUMAN-LIKE WRITING GUIDELINES (MANDATORY):**
        *   Clarity and Simplicity: Prioritize clear communication. Use straightforward language and sentence structures.
        *   Natural Vocabulary: Employ common, everyday words. Avoid overly complex, academic, or obscure vocabulary.
        *   Engaging Tone: Write in a way that captivates the reader. Make the prose flow naturally and conversationally.
        *   Avoid AI Tropes: Do not use overly formal language, excessive adverbs, or repetitive sentence beginnings that can make writing sound robotic.
    {optional_prompt_section}
        **OVERALL PLOT (For General Context):**
        {{full_plot}}
    
        **{plot_header}**
        {{focus_plot_description}}
    
        **CONTEXT (Previous Chapters, World Info, Project Structure):** 
        {{context}} 
        The 'CONTEXT' section above includes:
        1.  Previously generated chapters in this current batch - MOST IMPORTANT for maintaining direct narrative continuity
        2.  Relevant information from older chapters and the project's knowledge base.
        3.  The specific structural placement for *this* chapter (e.g., 'Current Structural Context for this Chapter: Act I > Stage 1 > Substage A').
        4.  The **FULL PROJECT STRUCTURE**, which lists all Acts, Stages, and Substages, along with any chapters currently linked to them. 
        Use this comprehensive context to understand the chapter's placement within the larger narrative arc and to ensure thematic and plot consistency.
    
        **NARRATIVE CONTINUITY REQUIREMENTS (CRITICAL):**
        1. Your chapter MUST coherently continue the storyline from any previous chapters in the current batch.
        2. Characters, settings, and plot elements should maintain precise consistency with previous chapters.
        3. Respect the chronological flow of events established in previous chapters.
        4. Avoid contradicting any events, dialogue or character decisions from previous chapters.
        5. Pick up naturally from the ending of the most recent chapter if applicable.

        **WRITING REQUIREMENTS (MANDATORY):**
        1.  Writing Style: {{writing_style}} (Apply this style METICULOUSLY to every sentence, keeping the Human-Like Writing Guidelines in mind.)
        2.  Style Guide: {{style_guide}} (Follow this guide EXACTLY, ensuring it aligns with the Human-Like Writing Guidelines.)
        3.  Additional Instructions: {{additional_instructions}} (Incorporate these precisely.)
        4.  Approximate Word Count Target: {{word_count_target}} words. Aim for this length naturally.
    
        **TASK:** Write Chapter {{chapter_number}}, focusing *primarily* on fulfilling the events and progression described in the **{plot_header.split('(')[0].strip()}** while maintaining perfect continuity with any previously generated chapters. Use the **OVERALL PLOT**, **PROJECT DESCRIPTION**, and the full **CONTEXT** (especially any previously generated chapters in this batch) to ensure consistency and depth, but the core action of this chapter should revolve around the plot segment/focus provided above.
    
        **FORMATTING (MANDATORY):**
        *   Use HTML `<p>` tags for paragraphs. Ensure proper paragraph separation.
        *   Inside `<p>` tags, format dialogue correctly (e.g., using quotation marks).
        *   Use HTML `<h3>***</h3>` for major scene breaks if appropriate.
    
        **CRITICAL RULES:**
        *   NEVER write the word "Codex".
        *   Follow the specified Writing Style and Style Guide WITHOUT FAIL.
        *   Stay EXACTLY true to the Plot/Setting focus provided for this chapter.
        *   Maintain PERFECT continuity with previous chapters in this generation batch.
        *   Incorporate ALL requirements.
        *   If specific sentence structures or endings are required, apply them CONSISTENTLY.
    
        Begin writing Chapter {{chapter_number}} immediately, starting directly with the first `<p>` tag. Do not include a chapter heading like "Chapter X: Title" or wrap the entire output in other tags like `<html>` or `<body>`.
        """
        human_template = "Write the chapter following all system instructions precisely, using HTML <p> tags for paragraphs."

        # Check for Anthropic to apply caching
        llm = state.get("llm")
        is_anthropic = False
        if llm and "Anthropic" in llm.__class__.__name__:
            is_anthropic = True

        if is_anthropic:
             sys_msg = SystemMessage(content=system_template)
             sys_msg.additional_kwargs["cache_control"] = {"type": "ephemeral"}
             return ChatPromptTemplate.from_messages(
                [sys_msg, ("human", human_template)]
             )

        return ChatPromptTemplate.from_messages(
            [("system", system_template), ("human", human_template)]
        )

    async def _generate_initial_chapter_node(
        self, state: ChapterGenerationState
    ) -> Dict[str, Any]:
        """Generates the initial draft of the chapter."""
        self.logger.debug(f"Node: Generating Initial Chapter {state['chapter_number']}")
        if state.get("error"):
            return {}  # Skip if error occurred previously

        try:
            llm = state["llm"]

            # Create the prompt using the updated helper, passing the whole state
            prompt = self._create_chapter_prompt(state=state)

            chain = prompt | llm | StrOutputParser()

            # Prepare the dictionary for invoking the chain
            # All necessary fields should be in the state and will be picked up by the prompt template
            # The ChatPromptTemplate will automatically pull the required fields from the state dict.
            invoke_dict = {
                **state,  # Pass the whole state, prompt template will pick what it needs
                "additional_instructions": state["instructions"].get(
                    "additionalInstructions", ""
                ),  # Ensure this specific key is available for the prompt
                "style_guide": state["instructions"].get("styleGuide", ""),
                "word_count_target": state["instructions"].get("wordCount", 0),
                "focus_plot_description": state.get("plot_segment")
                or state["full_plot"],  # Make sure this is directly available
            }

            # llm_response: BaseMessage = await chain.ainvoke(invoke_dict) # Old
            # The invoke_dict now directly comes from state, which ChatPromptTemplate can use.
            # Langchain templates will try to find the keys from the input dictionary.

            # The prompt expects: chapter_number, total_chapters, project_description, current_act_stage_info,
            # full_plot, focus_plot_description, context, writing_style, style_guide, additional_instructions, word_count_target

            # We need to ensure all these are top-level keys in the invoke_dict or accessible via state.
            # The state already has most of them. Let's ensure the prompt variables match state keys or are added to invoke_dict.

            # Create the specific dictionary for prompt variables
            prompt_variables = {
                "chapter_number": state["chapter_number"],
                "total_chapters": state["total_chapters"],
                "project_description": state.get("project_description"),
                "current_act_stage_info": state.get("current_act_stage_info"),
                "full_plot": state["full_plot"],
                "focus_plot_description": state.get("plot_segment")
                or state["full_plot"],
                "context": state["context"],
                "writing_style": state["writing_style"],
                "style_guide": state["instructions"].get("styleGuide", ""),
                "additional_instructions": state["instructions"].get(
                    "additionalInstructions", ""
                ),
                "word_count_target": state["instructions"].get("wordCount", 0),
            }

            llm_response_content = await chain.ainvoke(prompt_variables)
            # Since we are using StrOutputParser, the response is directly the string content.
            # If we were getting a BaseMessage, it would be: llm_response.content
            chapter_content = llm_response_content

            if not chapter_content or not chapter_content.strip():
                raise ValueError("LLM returned empty chapter content.")

            current_word_count = len(
                BeautifulSoup(chapter_content, "html.parser").get_text().split()
            )  # Estimate WC from text
            self.logger.info(
                f"Initial chapter generated. Word count: {current_word_count}"
            )

            return {
                "initial_chapter_content": chapter_content,
                "current_word_count": current_word_count,
                "target_word_count": state["instructions"].get(
                    "wordCount", 0
                ),  # Pass target along
                "last_llm_response": llm_response_content,  # Store the full AIMessage
            }
        except Exception as e:
            self.logger.error(
                f"Error in _generate_initial_chapter_node: {e}", exc_info=True
            )
            return {"error": f"Failed to generate initial chapter: {e}"}

    async def _extend_chapter_node(
        self, state: ChapterGenerationState
    ) -> Dict[str, Any]:
        """Extends the chapter if it's too short."""
        self.logger.debug(f"Node: Extending Chapter {state['chapter_number']}")
        if state.get("error"):
            return {}

        try:
            previous_content = (
                state.get("extended_chapter_content")
                or state["initial_chapter_content"]
            )
            current_word_count = state["current_word_count"]
            target_word_count = state["target_word_count"]
            words_to_add = target_word_count - current_word_count
            llm = state["llm"]
            last_llm_response = state.get("last_llm_response")

            if not last_llm_response:
                self.logger.error(
                    "Cannot extend chapter without a previous LLM response."
                )
                return {
                    "error": "Missing last_llm_response for extension.",
                    "extended_chapter_content": previous_content,
                    "current_word_count": current_word_count,
                }

            # Segmentation info for the human message
            plot_segment = state.get("plot_segment")
            focus_plot_description = (
                plot_segment if plot_segment else state["full_plot"]
            )
            plot_header_text = (
                "current chapter plot segment" if plot_segment else "plot/setting"
            )

            self.logger.info(
                f"Extending chapter. Current: {current_word_count}, Target: {target_word_count}, Need: {words_to_add}"
            )

            # Construct a human message to guide the LLM's continuation
            human_continuation_message_content = (
                f"Please continue writing the chapter. Add approximately {words_to_add} more words. "
                f"Remember to STRICTLY maintain the established writing style, tone, and focus on the {plot_header_text}: '{focus_plot_description}'. "
                f"Ensure the new text seamlessly follows the previous content.\n\n"
                f"**MANDATORY WRITING REQUIREMENTS (Apply these to the new text):**\n"
                f"1.  **Writing Style:** {state['writing_style']}\n"
                f"2.  **Style Guide:** {state['instructions'].get('styleGuide', '')}\n"
                f"3.  **Additional Instructions:** {state['instructions'].get('additionalInstructions', '')}\n\n"
                f"Output only the additional HTML paragraph(s), starting with a `<p>` tag. Ensure the new paragraphs adhere to the writing style."
            )
            human_message = HumanMessage(content=human_continuation_message_content)

            # Invoke the LLM with the history (last AI message) and the new human instruction
            # The LLM will treat this as a continued conversation
            messages_for_llm: List[BaseMessage] = [last_llm_response, human_message]

            current_llm_response: BaseMessage = await llm.ainvoke(messages_for_llm)
            extension_text = current_llm_response.content

            if not extension_text or not extension_text.strip():
                self.logger.warning(
                    "LLM returned empty extension text. Stopping extension."
                )
                return {
                    "extended_chapter_content": previous_content,
                    "current_word_count": current_word_count,
                    "last_llm_response": last_llm_response,  # Keep the previous successful response
                }

            # Combine HTML content
            full_content = (
                previous_content.strip() + "\n" + extension_text.strip()
            )  # Simple concatenation for HTML
            new_word_count = len(
                BeautifulSoup(full_content, "html.parser").get_text().split()
            )  # Estimate WC from text

            self.logger.info(f"Chapter extended. New word count: {new_word_count}")

            return {
                "extended_chapter_content": full_content,
                "current_word_count": new_word_count,
                "last_llm_response": current_llm_response,  # Store the new AIMessage from extension
            }

        except Exception as e:
            self.logger.error(f"Error in _extend_chapter_node: {e}", exc_info=True)
            # Return error but also the content we had before the failed extension attempt
            return {
                "error": f"Failed to extend chapter: {e}",
                "extended_chapter_content": state.get("extended_chapter_content")
                or state["initial_chapter_content"],
                "current_word_count": state["current_word_count"],
            }

    async def _generate_title_node(
        self, state: ChapterGenerationState
    ) -> Dict[str, Any]:
        """Generates the chapter title."""
        self.logger.debug(
            f"Node: Generating Title for Chapter {state['chapter_number']}"
        )
        if state.get("error"):
            return {}

        try:
            # Use the final content after potential extension
            final_content = (
                state.get("extended_chapter_content")
                or state["initial_chapter_content"]
            )
            chapter_number = state["chapter_number"]
            llm = state["check_llm"]  # Use check LLM for smaller tasks

            prompt = ChatPromptTemplate.from_template(
                """
            Analyze the following chapter content and generate a compelling, concise (2-6 words) title that captures its essence without spoilers. Maintain professional novel chapter naming conventions.

            Chapter Number: {chapter_number}
            Chapter Content Snippet:
            {chapter_content_snippet}

            Return ONLY the title text. Do not include "Chapter X:".
            """
            )

            chain = prompt | llm | StrOutputParser()

            # Provide a snippet to avoid large context for title gen
            content_snippet = (
                final_content[:1500] + "..."
                if len(final_content) > 1500
                else final_content
            )

            title_text = await chain.ainvoke(
                {
                    "chapter_number": chapter_number,
                    "chapter_content_snippet": content_snippet,
                }
            )

            cleaned_title = title_text.strip().replace('"', "")
            # Basic length check/truncation
            if len(cleaned_title) > 100:
                cleaned_title = cleaned_title[:97] + "..."
            if not cleaned_title:
                cleaned_title = f"Chapter {chapter_number} Untitled"

            final_title = f"Chapter {chapter_number}: {cleaned_title}"

            return {"chapter_title": final_title}
        except Exception as e:
            self.logger.error(f"Error in _generate_title_node: {e}", exc_info=True)
            # Return a default title if generation fails
            return {
                "error": f"Failed to generate title: {e}",
                "chapter_title": f"Chapter {state['chapter_number']}",
            }

    async def _extract_codex_items_node(
        self, state: ChapterGenerationState
    ) -> Dict[str, Any]:
        """Extracts new codex items from the chapter."""
        self.logger.debug(
            f"Node: Extracting Codex Items for Chapter {state['chapter_number']}"
        )
        if state.get("error"):
            return {}

        try:
            final_content = (
                state.get("extended_chapter_content")
                or state["initial_chapter_content"]
            )
            vector_store = state.get("vector_store")
            check_llm = state["check_llm"]
            user_id = state["user_id"]
            project_id = state["project_id"]

            # --- Comprehensive Existing Item Fetching ---
            existing_names = set()

            # 1. Fetch from Vector Store
            if vector_store:
                try:
                    # Fetch a larger number of existing items from the vector store
                    vector_store_docs = await vector_store.similarity_search(
                        query_text="*",  # Broad query
                        filter={
                            "type": {"$in": [t.value for t in CodexExtractionTypes]}
                        },
                        k=1000,  # Increased fetch limit
                    )
                    for doc in vector_store_docs:
                        if doc.metadata.get("name"):
                            existing_names.add(
                                self._normalize_name(doc.metadata.get("name"))
                            )
                except Exception as e:
                    self.logger.error(
                        f"Error fetching existing items from vector store: {e}"
                    )
            else:
                self.logger.error(
                    "Vector store is None. Cannot fetch existing items from vector store."
                )

            # 2. Fetch all items directly from the database for a definitive list
            try:
                db_items = await db_instance.get_all_codex_items(user_id, project_id)
                for item in db_items:
                    if item.get("name"):
                        existing_names.add(self._normalize_name(item.get("name")))
            except Exception as e:
                self.logger.error(f"Error fetching all codex items from DB: {e}")

            self.logger.debug(
                f"Compiled {len(existing_names)} unique existing codex item names for filtering."
            )

            parser = PydanticOutputParser(pydantic_object=CodexExtraction)

            # --- Strengthened Prompt ---
            prompt = ChatPromptTemplate.from_template(
                """
                Analyze the chapter content below. Identify NEW entities (characters, items, lore, factions, worldbuilding details) that are NOT in the 'Existing Item Names' list.

                IMPORTANT: You MUST NOT extract any entity if its name is present in the 'Existing Item Names' list. Check this list carefully. Even minor variations of existing names (e.g., different capitalization, punctuation) should be ignored.

                Valid Types: {valid_types}
                Valid Worldbuilding Subtypes: {valid_subtypes}

                CRITICAL RULE: Do NOT invent new types.
                - If you find a "Geography" or "Place", classify it as 'location'.
                - If you find a "History" or "Culture" detail, classify it as 'worldbuilding' with the appropriate subtype.

                For each NEW item:
                1. Extract its Name.
                2. Write a brief Description based ONLY on the chapter text.
                3. Assign its Type (from Valid Types).
                4. Assign a Subtype (from Valid Subtypes) ONLY if Type is 'worldbuilding'.

                Chapter Content:
                {chapter_content}

                Existing Item Names (IGNORE THESE - DO NOT EXTRACT ANY NAME FROM THIS LIST):
                {existing_names}

                Respond in JSON format matching the following schema:
                {format_instructions}

                If no new items are found, return an empty list for "new_items".
                """
            )

            all_new_items_raw = []
            try:
                llm_response_message = await (prompt | check_llm).ainvoke(
                    {
                        "chapter_content": final_content,
                        "existing_names": (
                            ", ".join(list(existing_names))
                            if existing_names
                            else "None"
                        ),
                        "valid_types": ", ".join(
                            [t.value for t in CodexExtractionTypes]
                        ),
                        "valid_subtypes": ", ".join(
                            [t.value for t in WorldbuildingSubtype]
                        ),
                        "format_instructions": parser.get_format_instructions(),
                    }
                )
                llm_output_text = llm_response_message.content
                if isinstance(llm_output_text, list):
                    llm_output_text = " ".join([str(item) for item in llm_output_text])
                elif not isinstance(llm_output_text, str):
                    llm_output_text = str(llm_output_text)

                if (
                    not llm_output_text
                    or not llm_output_text.strip()
                    or llm_output_text.strip().lower() == "null"
                ):
                    self.logger.info(
                        "LLM returned no new items. Defaulting to empty list."
                    )
                    validated_result = CodexExtraction(new_items=[])
                    all_new_items_raw = validated_result.new_items
                else:
                    # --- Robust "Row-Level" Parsing ---
                    # Instead of parsing the whole batch strictly, we parse JSON and then validate items one by one.
                    try:
                        # 1. Parse pure JSON first
                        json_parser = JsonOutputParser()
                        parsed_json = json_parser.parse(llm_output_text)
                        
                        raw_items_list = parsed_json.get("new_items", [])
                        valid_items_buffer = []

                        for raw_item in raw_items_list:
                            try:
                                # A. Attempt Automatic Correction of Common LLM Errors
                                if raw_item.get("type") == "geography":
                                    raw_item["type"] = "location"
                                    self.logger.info(f"Auto-corrected type 'geography' to 'location' for item '{raw_item.get('name')}'")

                                # B. Validate Individual Item
                                # We construct the Pydantic model for a single item
                                valid_item = ModelCodexItem(**raw_item)
                                valid_items_buffer.append(valid_item)
                            
                            except ValidationError as ve:
                                self.logger.warning(f"Skipping invalid item '{raw_item.get('name', 'UNKNOWN')}': {ve}")
                                continue # processing other items

                        all_new_items_raw = valid_items_buffer

                    except Exception as json_error:
                        self.logger.warning(
                            f"JSON parsing failed: {json_error}. Attempting fallback to fixing parser."
                        )
                        # Fallback: Try the fixing parser for the whole batch if basic JSON fails
                        fixing_parser = OutputFixingParser.from_llm(
                            parser=parser, llm=check_llm
                        )
                        validated_result = await fixing_parser.aparse(llm_output_text)
                        all_new_items_raw = validated_result.new_items

            except Exception as invoke_error:
                self.logger.error(
                    f"Codex extraction LLM call or parsing failed: {invoke_error}",
                    exc_info=True,
                )
                all_new_items_raw = []

            # --- Rigorous Post-Processing Validation ---
            final_validated_items = []
            llm_batch_unique_names = (
                set()
            )  # To de-duplicate items from the same LLM response

            for item in all_new_items_raw:
                if not item.name or not item.name.strip():
                    continue  # Skip items with no name

                normalized_name = self._normalize_name(item.name)

                # Check against comprehensive existing names AND duplicates in this batch
                if (
                    normalized_name in existing_names
                    or normalized_name in llm_batch_unique_names
                ):
                    self.logger.debug(
                        f"Filtering out duplicate item from LLM: '{item.name}' (Normalized: '{normalized_name}')"
                    )
                    continue

                # Validate Type
                try:
                    item_type_enum = CodexItemType(item.type)
                except ValueError:
                    self.logger.warning(
                        f"Invalid codex type '{item.type}' for item '{item.name}'. Skipping."
                    )
                    continue

                # Validate Subtype
                item_subtype_enum = None
                if item.subtype:
                    try:
                        item_subtype_enum = WorldbuildingSubtype(item.subtype)
                        if item_type_enum != CodexItemType.WORLDBUILDING:
                            self.logger.warning(
                                f"Subtype '{item.subtype}' provided for non-worldbuilding type '{item.type}' on item '{item.name}'. Ignoring subtype."
                            )
                            item_subtype_enum = None
                    except ValueError:
                        self.logger.warning(
                            f"Invalid worldbuilding subtype '{item.subtype}' for item '{item.name}'. Ignoring subtype."
                        )

                item_dict = {
                    "name": item.name.strip(),
                    "description": item.description.strip() if item.description else "",
                    "type": item_type_enum.value,
                    "subtype": item_subtype_enum.value if item_subtype_enum else None,
                }

                final_validated_items.append(item_dict)
                llm_batch_unique_names.add(
                    normalized_name
                )  # Add to batch set to prevent self-duplication

            self.logger.info(
                f"Extracted {len(final_validated_items)} unique new codex items after validation."
            )
            return {"new_codex_items": final_validated_items}

        except Exception as e:
            self.logger.error(f"Error in _extract_codex_items_node: {e}", exc_info=True)
            return {
                "error": f"Failed to extract codex items: {e}",
                "new_codex_items": [],  # Default to empty list on error
            }

    async def _validate_chapter_node(
        self, state: ChapterGenerationState
    ) -> Dict[str, Any]:
        """Validates the generated chapter."""
        self.logger.debug(f"Node: Validating Chapter {state['chapter_number']}")
        if state.get("error"):
            return {}

        try:
            final_content = (
                state.get("extended_chapter_content")
                or state["initial_chapter_content"]
            )
            instructions = state["instructions"]
            check_llm = state["check_llm"]
            vector_store = state["vector_store"]
            plot = state["plot"]

            # Fetch limited context for validation (e.g., plot + maybe previous chapter summary)
            # Re-using full context might be too much for validation LLM
            validation_context_docs = await vector_store.similarity_search(
                plot, k=3
            )  # Get relevant docs

            # Get previous chapters from the current batch for continuity validation
            previous_chapters_from_batch = instructions.get("previous_chapters", [])
            previous_chapters_context = ""
            if previous_chapters_from_batch:
                previous_chapters_context = (
                    "\nPreviously Generated Chapters In This Batch:\n"
                )
                for ch in previous_chapters_from_batch[
                    -1:
                ]:  # Just get the most recent one for validation
                    ch_num = ch.get("chapter_number", "?")
                    ch_title = ch.get("title", f"Chapter {ch_num}")
                    # Use a short excerpt from the end of the previous chapter
                    ch_content = ch.get("content", "")
                    if ch_content:
                        # Get the last 1000 chars for continuity check
                        ch_excerpt = (
                            ch_content[-1000:] if len(ch_content) > 1000 else ch_content
                        )
                        previous_chapters_context += f"Chapter {ch_num} ({ch_title}) Ending Excerpt:\n{ch_excerpt}\n\n"

            validation_context = (
                f"Plot: {plot}\n"
                f"Instructions: {json.dumps({k: v for k, v in instructions.items() if k != 'previous_chapters'})}\n"
                f"Chapter Number: {state['chapter_number']}\n"
                f"Total Chapters: {state['total_chapters']}\n"
                f"{previous_chapters_context}\n"
                f"Relevant Info:\n"
                + "\n".join(
                    [
                        f"- {doc.page_content[:200]}..."
                        for doc in validation_context_docs
                    ]
                )
            )

            parser = PydanticOutputParser(pydantic_object=ChapterValidation)
            fixing_parser = OutputFixingParser.from_llm(parser=parser, llm=check_llm)

            prompt = ChatPromptTemplate.from_template(
                """
            You are an expert editor evaluating a novel chapter based on provided criteria and context.

            **Chapter Content:**
            {chapter_content}

            **Evaluation Context & Instructions:**
            {validation_context}

            **Evaluation Criteria:**
            Plot Consistency, Character Development, Pacing, Dialogue Quality, Setting Description, Writing Style Adherence, Emotional Impact, Conflict/Tension, Theme Exploration, Grammar/Syntax, Narrative Continuity.

            **Special Focus on Narrative Continuity:**
            Closely evaluate how well this chapter maintains continuity with the previously generated chapters (if any):
            - Do events flow logically from the previous chapter?
            - Are characters, settings, and plot elements consistent?
            - Does the chapter avoid contradicting established facts?
            - Are open narrative threads from previous chapters appropriately continued?

            **Task:**
            1. Score each criterion from 1-10 (1=Poor, 10=Excellent) with a brief explanation.
            2. Assess overall validity (true/false).
            3. Provide an overall score (0-10).
            4. Give detailed scores/explanations specifically for Style Guide Adherence and Continuity.
            5. For the Continuity score in particular, focus on how well the chapter follows from previous chapters in the current batch.
            6. List specific Areas for Improvement.
            7. Write concise General Feedback.

            **Output Format (JSON):**
            {format_instructions}

            If you cannot perform the validation for any reason (e.g., content issues), return a JSON object with 'is_valid': false, 'overall_score': 0, and explain the reason in 'general_feedback'.
            """
            )

            # Define the LLM part of the chain separately
            llm_chain = prompt | check_llm

            # Invoke LLM first
            raw_llm_output = await llm_chain.ainvoke(
                {
                    "chapter_content": final_content,
                    "validation_context": validation_context,
                    "format_instructions": parser.get_format_instructions(),
                }
            )

            # Check if LLM output is usable
            content_to_parse = ""
            if raw_llm_output:
                if isinstance(raw_llm_output.content, list):
                    # Handle list content (e.g. from some models/configurations)
                    # Assuming list of strings or objects with 'text' attribute
                    parts = []
                    for part in raw_llm_output.content:
                        if isinstance(part, str):
                            parts.append(part)
                        elif hasattr(part, "text"):
                            parts.append(part.text)
                        else:
                            parts.append(str(part))
                    content_to_parse = "".join(parts)
                else:
                    content_to_parse = str(raw_llm_output.content)

            if not content_to_parse or not content_to_parse.strip():
                self.logger.error(
                    "Validation LLM returned empty or null content. Cannot parse."
                )
                # Return a default error structure
                return {
                    "error": "Validation LLM returned unusable output.",
                    "validity_check": {
                        "is_valid": False,
                        "overall_score": 0,
                        "general_feedback": "Validation could not be performed due to an LLM error or empty response.",
                        # Add other fields with default/null values if needed by downstream steps
                        "criteria_scores": {},
                        "style_guide_adherence_score": 0,
                        "style_guide_adherence_explanation": "N/A",
                        "continuity_score": 0,
                        "continuity_explanation": "N/A",
                        "areas_for_improvement": ["LLM validation failed."],
                    },
                }

            # Now try parsing with the fixing parser
            try:
                result = await fixing_parser.aparse(content_to_parse)
            except Exception as parse_error:
                self.logger.error(
                    f"Failed to parse validation output even with fixing parser: {parse_error}",
                    exc_info=True,
                )
                # Return a default error structure after parsing failure
                return {
                    "error": f"Failed to parse validation output: {parse_error}",
                    "validity_check": {
                        "is_valid": False,
                        "overall_score": 0,
                        "general_feedback": f"Validation output parsing failed: {parse_error}",
                        "criteria_scores": {},
                        "style_guide_adherence_score": 0,
                        "style_guide_adherence_explanation": "N/A",
                        "continuity_score": 0,
                        "continuity_explanation": "N/A",
                        "areas_for_improvement": ["Validation parsing failed."],
                    },
                }

            # Convert Pydantic result to simple dict for state
            validity_check_dict = {
                "is_valid": result.is_valid,
                "overall_score": result.overall_score,
                "criteria_scores": {
                    k: v.model_dump() for k, v in result.criteria_scores.items()
                },  # Store nested models as dicts
                "style_guide_adherence_score": result.style_guide_adherence.score,
                "style_guide_adherence_explanation": result.style_guide_adherence.explanation,
                "continuity_score": result.continuity.score,
                "continuity_explanation": result.continuity.explanation,
                "areas_for_improvement": result.areas_for_improvement,
                "general_feedback": result.general_feedback,
            }

            self.logger.info(
                f"Chapter validation completed. Overall Score: {result.overall_score}/10"
            )
            return {"validity_check": validity_check_dict}

        except (ValidationError, json.JSONDecodeError) as e:
            self.logger.error(f"Validation output parsing error: {e}", exc_info=True)
            return {"error": f"Failed to parse validation output: {e}"}
        except Exception as e:
            self.logger.error(f"Error in _validate_chapter_node: {e}", exc_info=True)
            return {"error": f"Failed during chapter validation: {e}"}

    async def _finalize_output_node(
        self, state: ChapterGenerationState
    ) -> Dict[str, Any]:
        """Assembles the final output dictionary."""
        self.logger.debug(
            f"Node: Finalizing Output for Chapter {state['chapter_number']}"
        )
        if state.get("error"):
            # Handle final error state
            return {"final_output": {"error": state["error"]}}

        final_content = (
            state.get("extended_chapter_content") or state["initial_chapter_content"]
        )

        final_output = {
            "content": final_content,
            "chapter_title": state.get(
                "chapter_title", f"Chapter {state['chapter_number']}"
            ),
            "new_codex_items": state.get("new_codex_items", []),
            "validity_check": state.get(
                "validity_check", {"error": "Validation skipped or failed."}
            ),
            "word_count": state.get("current_word_count", 0),
        }
        return {"final_output": final_output}

    # --- LangGraph Conditional Edges ---

    def _should_extend_chapter(self, state: ChapterGenerationState) -> str:
        """Determines if the chapter needs extension based on word count."""
        if state.get("error"):
            self.logger.warning("Error detected, skipping extension check.")
            return "proceed_to_title"  # Go to final steps even if error occurred

        current_wc = state["current_word_count"]
        target_wc = state["target_word_count"]
        initial_content = state["initial_chapter_content"]
        extended_content = state.get("extended_chapter_content")

        # Safety check: if initial generation failed, don't try to extend.
        if not initial_content:
            self.logger.error("Initial chapter generation failed, cannot extend.")
            return "proceed_to_title"

        # If we already tried extending and failed (e.g., empty response), don't loop infinitely.
        # Check if extension was attempted but didn't increase word count significantly.
        if (
            extended_content is not None
            and len(extended_content.split()) <= current_wc + 10
        ):  # Allowance for minor variations
            self.logger.warning(
                "Extension attempt did not significantly increase word count. Proceeding."
            )
            return "proceed_to_title"

        # Check if target word count is set and if current count is significantly lower
        # Use a threshold (e.g., 90%) to avoid unnecessary extensions for minor differences
        if target_wc > 0 and current_wc < target_wc * 0.9:
            self.logger.info(
                f"Word count {current_wc} is less than 90% of target {target_wc}. Extending."
            )
            return "extend_chapter"
        else:
            self.logger.info(
                f"Word count {current_wc} is sufficient (Target: {target_wc}). Proceeding."
            )
            return "proceed_to_title"

    # --- Graph Builder ---

    def _build_chapter_generation_graph(self) -> StateGraph:
        """Builds the LangGraph StateMachine for chapter generation."""
        graph = StateGraph(ChapterGenerationState)

        # Define nodes
        graph.add_node("construct_context", self._construct_context_node)
        graph.add_node("generate_initial_chapter", self._generate_initial_chapter_node)
        graph.add_node("extend_chapter", self._extend_chapter_node)
        graph.add_node("generate_title", self._generate_title_node)
        graph.add_node("extract_codex_items", self._extract_codex_items_node)
        graph.add_node("validate_chapter", self._validate_chapter_node)
        graph.add_node("finalize_output", self._finalize_output_node)

        # Define edges
        graph.set_entry_point("construct_context")
        graph.add_edge("construct_context", "generate_initial_chapter")

        # Conditional edge for extension
        graph.add_conditional_edges(
            "generate_initial_chapter",
            self._should_extend_chapter,
            {"extend_chapter": "extend_chapter", "proceed_to_title": "generate_title"},
        )
        # Loop back after extension attempt OR proceed if extension didn't work/isn't needed
        graph.add_conditional_edges(
            "extend_chapter",
            self._should_extend_chapter,  # Check again after extension
            {
                "extend_chapter": "extend_chapter",  # Loop if still too short and extension added words
                "proceed_to_title": "generate_title",  # Proceed if count is okay or extension failed
            },
        )

        graph.add_edge("generate_title", "extract_codex_items")
        graph.add_edge("extract_codex_items", "validate_chapter")
        graph.add_edge("validate_chapter", "finalize_output")
        graph.add_edge("finalize_output", END)

        # Compile the graph
        compiled_graph = graph.compile()
        self.logger.info("Chapter generation graph compiled.")
        return compiled_graph

    # --- Public Methods ---

    async def generate_chapter(
        self,
        chapter_number: int,
        plot: str,  # This is expected to be the FULL plot now
        writing_style: str,
        instructions: Dict[
            str, Any
        ],  # This now contains segment, full plot, total etc.
    ) -> Dict[str, Any]:
        """Generates a chapter using the LangGraph workflow."""
        self.logger.info(
            f"Starting chapter generation process for Chapter {chapter_number}..."
        )

        if not self.chapter_generation_graph:
            self.logger.error("Chapter generation graph is not initialized.")
            raise RuntimeError("AgentManager not properly initialized.")

        # Extract segmentation info from instructions, provide defaults if missing
        full_plot = instructions.get("full_plot", plot)  # Use passed plot as fallback
        plot_segment = instructions.get("plot_segment")  # Can be None
        total_chapters = instructions.get("total_chapters", 1)

        initial_state: ChapterGenerationState = {
            "user_id": self.user_id,
            "project_id": self.project_id,
            "chapter_number": chapter_number,
            "plot": plot,  # Keep the main 'plot' field as the FULL plot for context retrieval etc.
            "writing_style": writing_style,
            "instructions": instructions,  # Store original instructions
            # --- Add new fields ---
            "full_plot": full_plot,
            "plot_segment": plot_segment,
            "total_chapters": total_chapters,
            # --- End new fields ---
            "llm": self.llm,
            "check_llm": self.check_llm,
            "vector_store": self.vector_store,
            "summarize_chain": self.summarize_chain,
            # Initialize others to None/default
            "context": None,
            "initial_chapter_content": None,
            "extended_chapter_content": None,
            "current_word_count": 0,
            "target_word_count": instructions.get("wordCount", 0),
            "chapter_title": None,
            "new_codex_items": None,
            "validity_check": None,
            "error": None,
            "final_chapter_content": None,
            "final_output": None,
            "last_llm_response": None,  # Initialize new field
            "project_description": None,  # Added for project description
            "current_act_stage_info": None,  # Added for current act/stage/substage info
        }

        try:
            # Stream the graph execution (or use ainvoke for final result)
            # Using ainvoke for simplicity here, streaming requires more complex handling
            final_state = await self.chapter_generation_graph.ainvoke(initial_state)

            if final_state.get("error"):
                self.logger.error(
                    f"Chapter generation failed with error: {final_state['error']}"
                )
                # Return a structured error response
                return {
                    "error": final_state["error"],
                    "content": None,
                    "chapter_title": f"Chapter {chapter_number} (Failed)",
                    "new_codex_items": [],
                    "validity_check": None,
                }

            if not final_state.get("final_output"):
                raise ValueError("Graph execution finished without final output.")

            self.logger.info(
                f"Chapter generation process completed for Chapter {chapter_number}."
            )
            return final_state["final_output"]

        except Exception as e:
            self.logger.error(
                f"Error invoking chapter generation graph: {e}", exc_info=True
            )
            raise  # Re-raise the exception

    async def process_text_action(
        self,
        action: str,
        selected_text: str,
        full_chapter_content: str,
        custom_prompt: Optional[str] = None,
    ) -> str:
        """
        Processes a specific AI action (revise, extend, custom) on a selected piece of text
        within the context of the full chapter.
        """
        self.logger.info(
            f"Processing text action '{action}' for project {self.project_id}"
        )

        try:
            llm = self.llm  # Use the main LLM

            # Base prompt providing context and the core task
            system_prompt_template = """You are an expert editor. Your task is to modify a specific text selection based on a given action, while considering the full context of the chapter to maintain narrative and stylistic consistency.

            You must return ONLY the modified text selection, not the entire chapter. The returned text should seamlessly replace the original selection.

            **Full Chapter Context:**
            ---
            {full_chapter_content}
            ---

            **Original Text Selection to Modify:**
            ---
            {selected_text}
            ---

            **Action to Perform:** {action_description}
            """

            # Descriptions and specific instructions for each action
            action_descriptions = {
                "revise": "Revise the 'Original Text Selection' for clarity, impact, and flow. Improve the prose and word choice while preserving the core meaning and staying consistent with the chapter's context and style.",
                "extend": "Extend the 'Original Text Selection' by adding a few more sentences. The new content should logically and stylistically follow from the selection and fit within the chapter's context.",
                "custom": f"Apply the following instruction to the 'Original Text Selection': {custom_prompt}",
            }

            if action not in action_descriptions:
                raise ValueError(f"Invalid text action: {action}")

            # Construct the final prompt
            final_system_prompt = system_prompt_template.format(
                full_chapter_content=full_chapter_content,
                selected_text=selected_text,
                action_description=action_descriptions[action],
            )

            prompt = ChatPromptTemplate.from_messages(
                [
                    ("system", final_system_prompt),
                    (
                        "human",
                        "Please provide the modified text selection now, and nothing else.",
                    ),
                ]
            )

            chain = prompt | llm | StrOutputParser()

            # Invoke the chain
            modified_text = await chain.ainvoke({})  # All variables are in the formatted system prompt

            if not modified_text or not modified_text.strip():
                raise ValueError("LLM returned empty content for text action.")

            self.logger.info(f"Text action '{action}' completed successfully.")
            return modified_text.strip()

        except Exception as e:
            self.logger.error(f"Error in process_text_action: {e}", exc_info=True)
            raise

    async def save_validity_feedback(
        self, result: Dict[str, Any], chapter_number: int, chapter_id: str
    ):
        """Saves validity feedback to the database."""
        # This method remains largely the same, just ensure data format matches.
        try:
            # Get chapter title if needed (though often passed in result now)
            # chapter = await db_instance.get_chapter(chapter_id, self.user_id, self.project_id)
            chapter_title = (
                f"Chapter {chapter_number}"  # Get from state if possible later
            )

            # Adjust keys based on the structure from _validate_chapter_node
            await db_instance.save_validity_check(
                chapter_id=chapter_id,
                chapter_title=chapter_title,  # Or fetch dynamically
                is_valid=result["is_valid"],
                overall_score=result["overall_score"],
                general_feedback=result["general_feedback"],
                # Assuming criteria_scores is not directly saved, but specific scores are
                style_guide_adherence_score=result["style_guide_adherence_score"],
                style_guide_adherence_explanation=result[
                    "style_guide_adherence_explanation"
                ],
                continuity_score=result["continuity_score"],
                continuity_explanation=result["continuity_explanation"],
                areas_for_improvement=result[
                    "areas_for_improvement"
                ],  # Should be List[str]
                user_id=self.user_id,
                project_id=self.project_id,
            )
            self.logger.info(f"Validity feedback saved for chapter ID: {chapter_id}")
        except Exception as e:
            self.logger.error(
                f"Error saving validity feedback for chapter {chapter_id}: {e}",
                exc_info=True,
            )
            # Don't raise, as saving feedback is secondary to generation

    async def add_to_knowledge_base(
        self,
        content_type: str,
        content: str,
        metadata: Dict[str, Any],
        db_item_id: Optional[str] = None,
    ) -> Optional[str]:
        """Adds content to the vector store, optionally using a specific ID."""
        try:
            if not self.vector_store:
                self.logger.error("Vector store not initialized.")
                return None

            # Ensure mandatory metadata
            metadata["type"] = content_type
            metadata["user_id"] = self.user_id
            metadata["project_id"] = self.project_id
            metadata["created_at"] = datetime.now(timezone.utc).isoformat()
            # Add the source db_item_id to metadata for potential linking/debugging
            if db_item_id:
                metadata["id"] = db_item_id

            # Clean metadata (remove None values, ensure JSON serializable types)
            clean_metadata = {}
            for k, v in metadata.items():
                if v is not None:
                    if isinstance(v, (str, int, float, bool, list)):
                        clean_metadata[k] = v
                    elif isinstance(v, datetime):
                        clean_metadata[k] = v.isoformat()
                    # Add other serializable types if needed
                    else:
                        self.logger.warning(
                            f"Skipping non-serializable metadata key '{k}' of type {type(v)}"
                        )

            # Use add_texts for simplicity, VectorStore handles batching if implemented
            # Pass the db_item_id as the ID to use in the vector store
            ids = await self.vector_store.add_texts(
                [content], [clean_metadata], ids=[db_item_id] if db_item_id else None
            )
            if ids:
                embedding_id = ids[0]
                self.logger.debug(
                    f"Added '{content_type}' item to KB. Embedding/DB ID: {embedding_id}"
                )
                return embedding_id
            else:
                self.logger.error(
                    "Failed to add item to knowledge base, no ID returned."
                )
                return None

        except Exception as e:
            self.logger.error(f"Error adding to knowledge base: {e}", exc_info=True)
            return None  # Return None on error

    async def update_or_remove_from_knowledge_base(
        self,
        identifier: Union[str, Dict[str, str]],
        action: str,
        new_content: str = None,
        new_metadata: Dict[str, Any] = None,
    ):
        """Updates or removes an item from the vector store by embedding ID or item details."""
        if not self.vector_store:
            self.logger.error("Vector store not initialized.")
            return

        try:
            embedding_id = None
            original_db_id_from_identifier = (
                None  # To store item_id if identifier is a dict
            )

            if isinstance(identifier, str):
                embedding_id = identifier
            elif (
                isinstance(identifier, dict)
                and "item_id" in identifier
                and "item_type" in identifier
            ):
                original_db_id_from_identifier = identifier["item_id"]
                embedding_id = await self.vector_store.get_embedding_id(
                    item_id=original_db_id_from_identifier,
                    item_type=identifier["item_type"],
                )
                if not embedding_id:
                    self.logger.warning(
                        f"Could not find embedding ID for item: {identifier}. Cannot {action}."
                    )
                    return  # Return if embedding_id not found for dict identifier
            else:
                # This else block handles invalid identifier types
                raise ValueError(
                    "Invalid identifier provided. Must be embedding ID string or dict with item_id and item_type."
                )

            if action == "delete":
                if (
                    not embedding_id
                ):  # Should not happen if logic above is correct, but as a safeguard
                    self.logger.warning(
                        f"No embedding_id to delete for identifier {identifier}"
                    )
                    return
                await self.vector_store.delete_from_knowledge_base(embedding_id)
                self.logger.info(f"Deleted item with embedding ID: {embedding_id}")
            elif action == "update":
                if not embedding_id:  # Safeguard
                    self.logger.warning(
                        f"No embedding_id to update for identifier {identifier}"
                    )
                    return
                if new_content is None and new_metadata is None:
                    raise ValueError(
                        "Either new_content or new_metadata must be provided for update action"
                    )

                if new_metadata is None:
                    new_metadata = {}

                new_metadata["user_id"] = self.user_id
                new_metadata["project_id"] = self.project_id
                new_metadata["updated_at"] = datetime.now(timezone.utc).isoformat()

                # If the original identifier was a dict with item_id, preserve it as "id" in payload
                if original_db_id_from_identifier:
                    new_metadata["id"] = original_db_id_from_identifier
                elif (
                    isinstance(identifier, dict) and "item_id" in identifier
                ):  # Fallback just in case
                    new_metadata["id"] = identifier["item_id"]

                await self.vector_store.update_in_knowledge_base(
                    embedding_id, new_content, new_metadata
                )
                self.logger.info(f"Updated item with embedding ID: {embedding_id}")
            else:
                # This else block handles invalid action types
                raise ValueError(
                    f"Invalid action: {action}. Must be 'delete' or 'update'."
                )

        except Exception as e:
            self.logger.error(
                f"Error in update_or_remove_from_knowledge_base (ID: {identifier}, Action: {action}): {e}",
                exc_info=True,
            )
            # Avoid raising here to prevent breaking callers if KB update fails

    async def query_knowledge_base(
        self, query: str, chat_history: List[Dict[str, str]] = None
    ) -> Dict[str, Any]:  # Changed return type
        """Queries the knowledge base using RAG, acting as a chatbot."""
        # This method can remain largely the same, using the initialized components.
        if not self.vector_store or not self.model_settings:
            raise RuntimeError("AgentManager not initialized properly.")

        # self.logger.debug(
        #    f"Querying KB. Query: '{query}', History length: {len(chat_history) if chat_history else 0}"
        # )

        try:
            qa_llm = await self._get_llm(self.model_settings["knowledgeBaseQueryLLM"])

            # 1. Process History and Condense Question (for retrieval)
            standalone_question = query
            history_messages: List[BaseMessage] = []
            if chat_history:
                for msg in chat_history:
                    # Ensure content exists and is a string
                    content = msg.get("content", "")
                    if not isinstance(content, str):
                        content = str(content)  # Convert if not string

                    if msg.get("type") == "human":
                        history_messages.append(HumanMessage(content=content))
                    elif msg.get("type") == "ai":
                        history_messages.append(AIMessage(content=content))
                    else:
                        self.logger.warning(
                            f"Unknown chat history message type: {msg.get('type')}"
                        )

                if history_messages:
                    condense_prompt = ChatPromptTemplate.from_messages(
                        [
                            MessagesPlaceholder(variable_name="chat_history"),
                            (
                                "human",
                                "Given the chat history and the follow up question, rephrase the follow up question to be a standalone question, suitable for information retrieval.\n\nFollow Up Question: {question}\nStandalone question:",
                            ),
                        ]
                    )
                    # Use a cheaper/faster LLM potentially for condensing? Or same QA LLM.
                    condense_chain = condense_prompt | qa_llm | StrOutputParser()
                    try:
                        standalone_question = await condense_chain.ainvoke(
                            {"chat_history": history_messages, "question": query}
                        )
                        self.logger.debug(
                            f"Original Query: '{query}', Standalone Query for Retrieval: '{standalone_question}'"
                        )
                    except Exception as condense_err:
                        self.logger.warning(
                            f"Failed to condense question, using original query for retrieval: {condense_err}"
                        )
                        standalone_question = query  # Fallback

            # 2. Retrieve Documents using the standalone question
            retrieved_docs = await self.vector_store.similarity_search(
                standalone_question, k=5  # Limit sources for context window
            )

            sources_metadata = [doc.metadata for doc in retrieved_docs]
            self.logger.debug(f"Retrieved {len(retrieved_docs)} documents for context.")

            # If no docs found, maybe respond differently?
            # if not retrieved_docs:
            #     # Generate a response indicating no specific info found, maybe using history
            #     no_context_prompt = ...
            #     # return {"answer": no_context_answer, "sources": []}
            #     pass # For now, proceed and let LLM handle it

            # 3. Generate Answer using LLM with Full History and Context

            context_str = "\n\n---\n\n".join(
                [
                    # Include more metadata if helpful for the LLM?
                    f"Source {i+1} (Type: {doc.metadata.get('type', 'N/A')}, Name: {doc.metadata.get('name', 'Unknown')}):\n{doc.page_content}"
                    for i, doc in enumerate(retrieved_docs)
                ]
            )

            # New prompt incorporating history and context for conversational response
            answer_prompt = ChatPromptTemplate.from_messages(
                [
                    (
                        "system",
                        "You are a helpful AI assistant for this project, knowledgeable about its world and story. "
                        "Answer the user's query based on the provided conversation history and the relevant context passages retrieved from the project's knowledge base. "
                        "Be conversational and informative. If the context doesn't directly answer the question, state that clearly but try to provide related information from the context if possible. "
                        "Do not make up information not present in the context or history. Mentioning the source type or name (e.g., 'According to the notes on Character X...') can be helpful but is not mandatory."
                        "\n\nRelevant Context Passages:\n---\n{context}\n---",
                    ),
                    MessagesPlaceholder(
                        variable_name="chat_history"
                    ),  # Add history placeholder
                    ("human", "{question}"),  # The user's latest question
                ]
            )

            answer_chain = answer_prompt | qa_llm | StrOutputParser()

            # Invoke with history, context, and the original query
            llm_answer = await answer_chain.ainvoke(
                {
                    "context": context_str,
                    "chat_history": history_messages,  # Pass the processed history objects
                    "question": query,  # Use the original query here
                }
            )

            # self.logger.debug(f"LLM Answer: {llm_answer[:100]}...")

            return {"answer": llm_answer, "sources": sources_metadata}

        except Exception as e:
            self.logger.error(f"Error in query_knowledge_base: {e}", exc_info=True)
            # Return a structured error
            return {
                "answer": f"An error occurred while processing your request: {e}",
                "sources": [],
            }

    # --- Other Methods (Keep as needed, ensure they use async/await and initialized components) ---

    async def generate_codex_item(
        self, codex_type: str, subtype: Optional[str], description: str
    ) -> Dict[str, str]:
        """Generates details for a new codex item based on a description."""
        # This can remain a separate utility, doesn't need full graph
        self.logger.debug(
            f"Generating codex item: Type={codex_type}, Subtype={subtype}"
        )
        try:
            parser = PydanticOutputParser(
                pydantic_object=ModelCodexItem
            )  # Reformatted line
            llm = await self._get_llm(
                self.model_settings["extractionLLM"]
            )  # Use extraction LLM
            fixing_parser = OutputFixingParser.from_llm(parser=parser, llm=llm)

            # Fetch relevant context (existing items, chapters)
            existing_items = await db_instance.get_all_codex_items(
                self.user_id, self.project_id
            )
            existing_items_str = json.dumps(
                [
                    {
                        "name": i.get("name"),
                        "type": i.get("type"),
                        "desc_snippet": i.get("description", "")[:50] + "...",
                    }
                    for i in existing_items[:20]
                ],
                indent=2,
            )  # Snippets only

            relevant_chapters = await self.vector_store.similarity_search(
                query_text=description, filter={"type": "chapter"}, k=5
            )
            chapter_context = "\n\n".join(
                f"Ch {doc.metadata.get('chapter_number', 'N/A')} Snippet: {doc.page_content[:300]}..."
                for doc in relevant_chapters
            )

            prompt = ChatPromptTemplate.from_template(
                """
            Generate a name and detailed description for a new codex item.

            **Item Specifications:**
            *   Type: {codex_type}
            *   Subtype: {subtype}
            *   Initial Description Idea: {description}

            **Context:**
            *   Relevant Chapter Snippets: {chapter_context}
            *   Some Existing Codex Items: {existing_codex_items}

            **Task:**
            1.  Create a concise, evocative Name.
            2.  Write a comprehensive Description, expanding on the initial idea with depth and vivid details, consistent with the context.
            3.  Ensure consistency with the specified type/subtype and existing items/chapters. DO NOT contradict existing info.
            4.  Follow specific guidelines for the type (Character: appearance, personality, role; Item: origin, properties; Lore: date, impact; Worldbuilding: details based on subtype).

            **Output Format (JSON):**
            {format_instructions}
            """
            )

            chain = prompt | llm | fixing_parser

            result = await chain.ainvoke(
                {
                    "codex_type": codex_type,
                    "subtype": subtype or "N/A",
                    "description": description,
                    "existing_codex_items": existing_items_str,
                    "chapter_context": chapter_context,
                    "format_instructions": parser.get_format_instructions(),
                }
            )

            return {"name": result.name, "description": result.description}

        except Exception as e:
            self.logger.error(f"Error generating codex item: {e}", exc_info=True)
            return {"name": "Error", "description": f"Failed to generate: {e}"}

    async def analyze_character_relationships(
        self, characters: List[Dict[str, Any]]
    ) -> List[RelationshipAnalysis]:
        """Analyzes and potentially saves relationships between the provided characters based on project context."""
        self.logger.info(f"Analyzing relationships for {len(characters)} characters.")
        if not characters or len(characters) < 2:
            self.logger.warning(
                "Need at least two characters to analyze relationships."
            )
            return []

        try:
            # Get story context (consider summarizing or selecting relevant parts if too large)
            all_chapters_data = await db_instance.get_all_chapters(
                self.user_id, self.project_id
            )
            context_content = "\n\n".join(
                [
                    f"Chapter {c.get('chapter_number', 'N/A')}: {c.get('content', '')[:2000]}..."
                    for c in all_chapters_data
                ]
            )  # Snippets

            context_tokens = self.estimate_token_count(context_content)
            max_context_tokens = (
                self.MAX_INPUT_TOKENS // 3
            )  # Allocate portion for context
            if context_tokens > max_context_tokens:
                self.logger.warning(
                    f"Relationship analysis context too large ({context_tokens} tokens). Summarizing."
                )
                # Summarize context (or use vector search for relevance)
                docs_to_summarize = [Document(page_content=context_content)]
                # Use ainvoke with a dictionary input
                summary_result = await self.summarize_chain.ainvoke(
                    {"input_documents": docs_to_summarize}
                )
                # Access the result
                context_content = summary_result.get(
                    "output_text", "Summary unavailable"
                )

            parser = PydanticOutputParser(pydantic_object=RelationshipAnalysisList)
            # Use check_llm or extractionLLM as configured
            relationship_llm = await self._get_llm(self.model_settings["extractionLLM"])
            fixing_parser = OutputFixingParser.from_llm(
                parser=parser, llm=relationship_llm
            )

            prompt = ChatPromptTemplate.from_template(
                """
            Analyze the relationships between the provided characters based on the story context. Focus ONLY on pairs involving the listed characters.

            Characters to Analyze:
            {characters_json}

            Story Context (Summaries/Snippets):
            {context}

            For each significant relationship pair found between the specified characters:
            1. Identify the names of both characters (character1, character2).
            2. Determine the relationship type (e.g., friend, enemy, rival, family, mentor, romantic interest, etc.).
            3. Provide a concise description summarizing their interactions and feelings towards each other based *only* on the context.

            Return ONLY the relationships between pairs of characters from the provided list. Do not infer relationships not present in the context.

            Format your response as JSON:
            {format_instructions}
            """
            )

            chain = (
                prompt | relationship_llm | fixing_parser
            )  # Use fixing parser as fallback

            character_list_json = json.dumps(
                [{"id": c.get("id"), "name": c.get("name")} for c in characters],
                indent=2,
            )

            result = await chain.ainvoke(
                {
                    "characters_json": character_list_json,
                    "context": context_content,
                    "format_instructions": parser.get_format_instructions(),
                }
            )

            saved_relationships_output = []
            processed_pairs = set()  # To avoid duplicate db entries for A-B and B-A

            if hasattr(result, "relationships"):
                for rel in result.relationships:
                    # Find the actual character dicts from the input list
                    char1_dict = next(
                        (c for c in characters if c.get("name") == rel.character1), None
                    )
                    char2_dict = next(
                        (c for c in characters if c.get("name") == rel.character2), None
                    )

                    if char1_dict and char2_dict:
                        # Ensure pair uniqueness (regardless of order)
                        pair_key = tuple(sorted([char1_dict["id"], char2_dict["id"]]))
                        if pair_key not in processed_pairs:
                            processed_pairs.add(pair_key)
                            try:
                                # Save to database (assuming create_character_relationship exists)
                                # Note: Original method saved to db, decide if that's desired here or just return analysis
                                relationship_id = (
                                    await db_instance.create_character_relationship(
                                        character_id=char1_dict["id"],
                                        related_character_id=char2_dict["id"],
                                        relationship_type=rel.relationship_type,
                                        project_id=self.project_id,
                                        description=rel.description,
                                    )
                                )
                                self.logger.debug(
                                    f"Saved relationship between {rel.character1} and {rel.character2} (ID: {relationship_id})"
                                )

                                # Optionally add relationship info to knowledge base
                                relationship_content = f"Relationship between {rel.character1} and {rel.character2}: {rel.relationship_type}. {rel.description}"
                                await self.add_to_knowledge_base(
                                    "relationship",
                                    relationship_content,
                                    {
                                        "name": f"{rel.character1}-{rel.character2} relationship",
                                        "type": "relationship",
                                        "relationship_id": relationship_id,  # Link to DB entry
                                        "character1_id": char1_dict["id"],
                                        "character2_id": char2_dict["id"],
                                    },
                                )
                                # Add the Pydantic model directly to the output list
                                saved_relationships_output.append(rel)

                            except Exception as db_error:
                                self.logger.error(
                                    f"Failed to save relationship between {rel.character1} and {rel.character2}: {db_error}"
                                )
                                # Continue processing other relationships
            else:
                self.logger.warning(
                    "Relationship analysis LLM call returned no 'relationships' field."
                )

            return saved_relationships_output

        except Exception as e:
            self.logger.error(
                f"Error analyzing character relationships: {e}", exc_info=True
            )
            raise  # Re-raise the error

    async def analyze_character_journey(self, character_id: str) -> Optional[str]:
        """Analyzes the manuscript content to generate a plausible backstory for the character."""
        try:
            self.logger.info(
                f"Starting backstory generation from manuscript for: {character_id}"
            )

            # 1. Fetch Character Details (Name)
            character = await db_instance.get_codex_item_by_id(
                character_id, self.user_id, self.project_id
            )
            if not character:
                self.logger.error(
                    f"Character {character_id} not found for backstory generation."
                )
                return None
            if character.get("type") != CodexItemType.CHARACTER.value:
                self.logger.error(f"Item {character_id} is not a character.")
                return None

            character_name = character.get("name", "Unknown Character")
            self.logger.debug(f"Generating backstory for character: {character_name}")

            # 2. Fetch All Chapter Content
            all_chapters = await db_instance.get_all_chapters(
                self.user_id, self.project_id
            )
            if not all_chapters:
                self.logger.warning(
                    f"No chapters found for project {self.project_id}. Cannot generate backstory."
                )
                return "No manuscript content found to analyze."

            # Sort chapters (optional but good practice)
            all_chapters.sort(key=lambda x: x.get("chapter_number", 0))
            manuscript_text = "\n\n---\n\n".join(
                [ch.get("content", "") for ch in all_chapters]
            )
            self.logger.debug(
                f"Combined manuscript text length: {len(manuscript_text)}"
            )

            # 3. Define the Backstory Generation Prompt
            prompt_template = ChatPromptTemplate.from_messages(
                [
                    (
                        "system",
                        "You are an AI assistant skilled at analyzing fictional narratives and inferring character history. "
                        "Read the provided manuscript content carefully. Your task is to identify the character '{character_name}' and generate a plausible backstory for them based *only* on their actions, dialogue, relationships, and any hints mentioned within the text. "
                        "Focus on creating a concise, compelling narrative that explains their origins, motivations, or key past events implied by the story. "
                        "Do not invent information not supported by the text. If the text provides little information about the character's past, state that clearly and generate a minimal backstory based only on available clues.",
                    ),
                    (
                        "human",
                        "Character Name: {character_name}\n\nManuscript Content:\n{manuscript_text}",
                    ),
                ]
            )

            # 4. Prepare the LLM chain (using main llm for generation)
            chain = prompt_template | self.llm | StrOutputParser()

            # 5. Invoke LLM
            self.logger.debug(
                f"Invoking LLM for backstory generation of {character_name}"
            )
            # Basic truncation if manuscript is extremely long - more sophisticated handling might be needed
            max_llm_input = self.MAX_INPUT_TOKENS - 1000  # Reserve tokens for prompt
            if self.estimate_token_count(manuscript_text) > max_llm_input:
                self.logger.warning(
                    f"Manuscript text potentially too long ({self.estimate_token_count(manuscript_text)} tokens). Truncating for LLM."
                )
                # Simple truncation - find better way later if needed
                ratio = max_llm_input / self.estimate_token_count(manuscript_text)
                truncated_manuscript = manuscript_text[
                    : int(len(manuscript_text) * ratio)
                ]
            else:
                truncated_manuscript = manuscript_text

            generated_backstory = await chain.ainvoke(
                {
                    "character_name": character_name,
                    "manuscript_text": truncated_manuscript,
                }
            )
            self.logger.debug(
                f"LLM Generated Backstory: {generated_backstory[:150]}..."
            )

            # 6. Return the generated backstory
            return generated_backstory

        except Exception as e:
            self.logger.error(
                f"Error generating backstory for {character_id}: {str(e)}",
                exc_info=True,
            )
            return f"An error occurred during backstory generation: {str(e)}"  # Return error message

    async def analyze_unprocessed_chapter_locations(self) -> List[Dict[str, Any]]:
        """
        Analyzes unprocessed chapters for locations, adds them to the database and knowledge base,
        and marks chapters as processed. This version includes batching and deduplication.
        """
        self.logger.info(
            f"Starting location analysis for unprocessed chapters in project {self.project_id}"
        )
        try:
            # 1. Fetch existing locations for context and deduplication
            existing_locations_data = await db_instance.get_locations(
                self.user_id, self.project_id
            )
            existing_location_names = {
                self._normalize_name(loc["name"]) for loc in existing_locations_data
            }
            self.logger.debug(
                f"Found {len(existing_location_names)} existing locations."
            )

            # 2. Fetch all unprocessed chapter content
            unprocessed_chapters = (
                await db_instance.get_latest_unprocessed_chapter_content(
                    self.project_id, self.user_id, "locations_analyzed"
                )
            )

            if not unprocessed_chapters:
                self.logger.info("No unprocessed chapters found for location analysis.")
                return []

            self.logger.info(
                f"Found {len(unprocessed_chapters)} unprocessed chapters for location analysis."
            )

            # 3. Process chapters in batches
            batch_size = 10  # Process 10 chapters at a time
            all_new_locations = []

            for i in range(0, len(unprocessed_chapters), batch_size):
                batch = unprocessed_chapters[i : i + batch_size]
                self.logger.info(f"Processing batch {i//batch_size + 1} of chapters.")

                batch_content = "\n\n".join(
                    [
                        f"--- Chapter Content ---\n{chapter['content']}"
                        for chapter in batch
                    ]
                )

                # 4. Create prompt and invoke LLM
                prompt = ChatPromptTemplate.from_messages(
                    [
                        (
                            "system",
                            """You are an expert at analyzing story content to identify and extract key locations.
Your task is to identify potential locations from the provided chapter content.
- A location should be a specific place, like a city, building, forest, or room.
- Do not extract general concepts like 'the past' or 'her memories'.
- Compare against the list of existing locations provided and only output NEW locations that are not on the list.
- For each new location, provide a name and a detailed description based on the text.
- If no new locations are found, return an empty list.""",
                        ),
                        (
                            "human",
                            """Here is the list of existing locations in the project for your reference:
{existing_locations}

Here is the content of the latest chapters to analyze:
{chapter_content}

Based on the chapter content, identify and extract any new locations.
Your response should be a JSON object with a single key "locations" which is a list of objects, where each object has "name" and "description" keys.""",
                        ),
                    ]
                )

                llm = await self._get_llm(self.model_settings["extractionLLM"])
                chain = prompt | llm | JsonOutputParser()

                # 5. Invoke the analysis chain
                try:
                    analysis_result = await chain.ainvoke(
                        {
                            "existing_locations": (
                                ", ".join(sorted(existing_location_names))
                                if existing_location_names
                                else "None"
                            ),
                            "chapter_content": batch_content,
                        }
                    )
                    batch_new_locations = analysis_result.get("locations", [])
                except Exception as e:
                    self.logger.error(f"Error during location analysis LLM call: {e}")
                    continue  # Move to the next batch

                # 6. Process and save new locations from the batch
                for location in batch_new_locations:
                    normalized_name = self._normalize_name(location.get("name", ""))
                    if (
                        not normalized_name
                        or normalized_name in existing_location_names
                    ):
                        continue  # Skip empty or duplicate names

                    new_loc_id = await db_instance.create_location(
                        name=location["name"],
                        description=location.get("description", ""),
                        coordinates=None,
                        user_id=self.user_id,
                        project_id=self.project_id,
                    )
                    self.logger.info(
                        f"Created new location: {location['name']} ({new_loc_id})"
                    )

                    # Add to knowledge base
                    await self.add_to_knowledge_base(
                        content_type="location",
                        content=f"Location Name: {location['name']}\nDescription: {location.get('description', '')}",
                        metadata={
                            "name": location["name"],
                            "source": "Chapter Analysis",
                        },
                        db_item_id=new_loc_id,
                    )

                    all_new_locations.append({"id": new_loc_id, **location})
                    existing_location_names.add(
                        normalized_name
                    )  # Add to set to prevent duplicates within the same run

                # 7. Mark chapters in the batch as processed
                for chapter in batch:
                    await db_instance.mark_chapter_processed(
                        chapter["id"], self.user_id, "locations_analyzed"
                    )
                self.logger.info(
                    f"Marked {len(batch)} chapters as processed for location analysis."
                )

            self.logger.info(
                f"Location analysis complete. Found {len(all_new_locations)} new locations."
            )
            return all_new_locations

        except Exception as e:
            self.logger.error(
                f"An error occurred during location analysis: {e}", exc_info=True
            )
            return []

    async def analyze_unprocessed_chapter_events(self) -> List[Dict[str, Any]]:
        """
        Analyzes unprocessed chapters for events, adds them to the database and knowledge base,
        and marks chapters as processed. This version includes batching and deduplication.
        """
        self.logger.info(
            f"Starting event analysis for unprocessed chapters in project {self.project_id}"
        )
        try:
            # 1. Fetch existing events for context and deduplication
            existing_events_data = await db_instance.get_events(
                self.project_id, self.user_id, limit=None  # Fetch all for dedupe
            )
            existing_event_titles = {
                self._normalize_name(event["title"]) for event in existing_events_data
            }
            self.logger.debug(f"Found {len(existing_event_titles)} existing events.")

            # 2. Fetch all unprocessed chapter content
            unprocessed_chapters = (
                await db_instance.get_latest_unprocessed_chapter_content(
                    self.project_id, self.user_id, "events_analyzed"
                )
            )

            if not unprocessed_chapters:
                self.logger.info("No unprocessed chapters found for event analysis.")
                return []

            self.logger.info(
                f"Found {len(unprocessed_chapters)} unprocessed chapters for event analysis."
            )

            # 3. Process chapters in batches
            batch_size = 10
            all_new_events = []

            for i in range(0, len(unprocessed_chapters), batch_size):
                batch = unprocessed_chapters[i : i + batch_size]
                self.logger.info(
                    f"Processing event batch {i//batch_size + 1} of chapters."
                )

                batch_content = "\n\n---\n\n".join(
                    [
                        f"--- Chapter Content ---\n{chapter['content']}"
                        for chapter in batch
                    ]
                )

                # 4. Create prompt and invoke LLM
                prompt = ChatPromptTemplate.from_messages(
                    [
                        (
                            "system",
                            """You are an expert at analyzing story content to identify and extract key events.
Your task is to identify potential events from the provided chapter content.
- An event is a significant occurrence or happening in the story that affects characters or plot.
- Do not extract general themes or ongoing states. Focus on specific actions or turning points.
- Compare against the list of existing events provided and only output NEW events that are not on the list.
- For each new event, provide a concise title and a detailed description based on the text.
- If no new events are found, return an empty list.""",
                        ),
                        (
                            "human",
                            """Here is the list of existing events in the project for your reference:
{existing_events}

Here is the content of the latest chapters to analyze:
{chapter_content}

Based on the chapter content, identify and extract any new events.
Your response should be a JSON object with a single key "events" which is a list of objects, where each object has "title" and "description" keys.""",
                        ),
                    ]
                )

                llm = await self._get_llm(self.model_settings["extractionLLM"])
                chain = prompt | llm | JsonOutputParser()

                # 5. Invoke the analysis chain
                try:
                    analysis_result = await chain.ainvoke(
                        {
                            "existing_events": (
                                ", ".join(sorted(existing_event_titles))
                                if existing_event_titles
                                else "None"
                            ),
                            "chapter_content": batch_content,
                        }
                    )
                    batch_new_events = analysis_result.get("events", [])
                except Exception as e:
                    self.logger.error(f"Error during event analysis LLM call: {e}")
                    continue

                # 6. Process and save new events from the batch
                for event in batch_new_events:
                    normalized_title = self._normalize_name(event.get("title", ""))
                    if (
                        not normalized_title
                        or normalized_title in existing_event_titles
                    ):
                        continue

                    db_event_id = await db_instance.create_event(
                        title=event["title"],
                        description=event.get("description", ""),
                        date=datetime.now(timezone.utc),  # Placeholder date
                        project_id=self.project_id,
                        user_id=self.user_id,
                    )
                    self.logger.info(
                        f"Created new event: {event['title']} ({db_event_id})"
                    )

                    await self.add_to_knowledge_base(
                        content_type="event",
                        content=f"Event Title: {event['title']}\nDescription: {event.get('description', '')}",
                        metadata={
                            "title": event["title"],
                            "source": "Chapter Analysis",
                        },
                        db_item_id=db_event_id,
                    )

                    all_new_events.append({"id": db_event_id, **event})
                    existing_event_titles.add(normalized_title)

                # 7. Mark chapters in the batch as processed
                for chapter in batch:
                    await db_instance.mark_chapter_processed(
                        chapter["id"], self.user_id, "events_analyzed"
                    )
                self.logger.info(
                    f"Marked {len(batch)} chapters as processed for event analysis."
                )

            self.logger.info(
                f"Event analysis complete. Found {len(all_new_events)} new events."
            )
            return all_new_events

        except Exception as e:
            self.logger.error(
                f"An error occurred during event analysis: {e}", exc_info=True
            )
            return []

    async def analyze_event_connections(self) -> List[EventConnectionBase]:
        """Analyzes connections between events in the project, processing in batches."""
        self.logger.info(
            f"Starting event connection analysis for project {self.project_id}"
        )
        try:
            events_data = await db_instance.get_events(self.project_id, self.user_id)
            if len(events_data) < 2:
                self.logger.info("Not enough events to analyze for connections.")
                return []

            self.logger.info(
                f"Found {len(events_data)} events to analyze for connections."
            )

            # Get existing connections to avoid re-creating them
            existing_connections_data = await db_instance.get_event_connections(
                self.project_id, self.user_id
            )
            existing_pairs = {
                tuple(sorted((c["event1_id"], c["event2_id"])))
                for c in existing_connections_data
            }

            all_new_connections = []
            batch_size = 25  # Number of events to consider in each batch

            # Create combinations of events to check
            event_pairs = list(combinations(events_data, 2))

            # Get all chapters for context retrieval
            chapters = await db_instance.get_all_chapters(self.user_id, self.project_id)
            chapter_content_map = {
                chapter["id"]: chapter["content"] for chapter in chapters
            }

            for i in range(0, len(event_pairs), batch_size):
                batch = event_pairs[i : i + batch_size]
                self.logger.info(
                    f"Processing batch {i//batch_size + 1} of event pairs."
                )

                # Filter out pairs that already have a connection
                batch_to_analyze = [
                    pair
                    for pair in batch
                    if tuple(sorted((pair[0]["id"], pair[1]["id"])))
                    not in existing_pairs
                ]

                if not batch_to_analyze:
                    self.logger.debug(
                        "All pairs in this batch already have connections."
                    )
                    continue

                # Enhanced formatting with RAG context
                formatted_pairs_with_context = []

                for idx, pair in enumerate(batch_to_analyze):
                    event1 = pair[0]
                    event2 = pair[1]

                    # Retrieve chapter context for each event
                    event1_context = await self._get_event_chapter_context(
                        event1["title"], chapter_content_map
                    )
                    event2_context = await self._get_event_chapter_context(
                        event2["title"], chapter_content_map
                    )

                    pair_info = f"Event Pair {idx+1}:\n"
                    pair_info += f"- Event 1: {event1['title']} (ID: {event1['id']})\n"
                    pair_info += f"  Description: {event1.get('description', 'No description')}\n"

                    if event1_context:
                        pair_info += f"  Context from chapters: {event1_context}\n"

                    pair_info += f"- Event 2: {event2['title']} (ID: {event2['id']})\n"
                    pair_info += f"  Description: {event2.get('description', 'No description')}\n"

                    if event2_context:
                        pair_info += f"  Context from chapters: {event2_context}\n"

                    formatted_pairs_with_context.append(pair_info)

                formatted_pairs = "\n\n".join(formatted_pairs_with_context)

                prompt = ChatPromptTemplate.from_template(
                    """You are a master storyteller and plot analyst. Your task is to identify meaningful connections between pairs of events.
A connection could be causal (one event causes another), thematic (they share a common theme), or consequential (one event is a consequence of another).

Analyze the following pairs of events and identify if a meaningful connection exists.
For each pair that is connected, explain the nature and impact of this connection.
Use the provided context from chapters when available to help determine connections.

Event Pairs to Analyze:
{event_pairs}

Respond in JSON format with a single key "connections", which is a list of objects.
Each object must contain:
- "event1_id": The ID of the first event.
- "event2_id": The ID of the second event.
- "connection_type": A brief type for the connection (e.g., "Causal", "Thematic", "Consequence").
- "description": A detailed explanation of how the events are connected.
- "impact": The significance of this connection to the overall plot or characters.

If a pair is not connected, do not include it in your response. If no connections are found in any of the pairs, return an empty list.
"""
                )

                llm = await self._get_llm(self.model_settings["extractionLLM"])
                chain = prompt | llm | JsonOutputParser()

                try:
                    result = await chain.ainvoke({"event_pairs": formatted_pairs})
                    new_connections = result.get("connections", [])
                except Exception as e:
                    self.logger.error(
                        f"Error during event connection analysis LLM call: {e}"
                    )
                    continue

                for conn in new_connections:
                    event1_id = conn.get("event1_id")
                    event2_id = conn.get("event2_id")

                    if not event1_id or not event2_id:
                        continue

                    # Avoid duplicates within the run
                    pair_key = tuple(sorted((event1_id, event2_id)))
                    if pair_key in existing_pairs:
                        continue

                    db_connection_id = await db_instance.create_event_connection(
                        event1_id=event1_id,
                        event2_id=event2_id,
                        connection_type=conn.get("connection_type", "Undefined"),
                        description=conn.get("description", ""),
                        impact=conn.get("impact", ""),
                        project_id=self.project_id,
                        user_id=self.user_id,
                    )
                    self.logger.info(
                        f"Created connection between events {event1_id} and {event2_id}."
                    )
                    conn["id"] = db_connection_id
                    all_new_connections.append(EventConnectionBase(**conn))
                    existing_pairs.add(pair_key)

            self.logger.info(
                f"Event connection analysis complete. Found {len(all_new_connections)} new connections."
            )
            return all_new_connections
        except Exception as e:
            self.logger.error(
                f"An error occurred during event connection analysis: {e}",
                exc_info=True,
            )
            return []

    async def _get_event_chapter_context(
        self, event_title: str, chapter_content_map: Dict[str, str]
    ) -> str:
        """
        Retrieves relevant context from chapters that mention this event.
        Returns a concise summary of the mentions.
        """
        relevant_excerpts = []

        # Look for event title mentions in all chapters
        for chapter_id, content in chapter_content_map.items():
            if event_title in content:
                # Find the paragraph containing the event mention
                paragraphs = content.split("\n\n")
                for paragraph in paragraphs:
                    if event_title in paragraph:
                        # Truncate long paragraphs
                        if len(paragraph) > 300:
                            start_pos = max(0, paragraph.find(event_title) - 100)
                            excerpt = paragraph[start_pos : start_pos + 300] + "..."
                        else:
                            excerpt = paragraph

                        relevant_excerpts.append(excerpt)
                        break  # Just get one mention per chapter

        # Limit the total context length
        if not relevant_excerpts:
            return ""

        # Join up to 3 excerpts
        if len(relevant_excerpts) > 3:
            return (
                "\n".join(relevant_excerpts[:3])
                + f" (and {len(relevant_excerpts) - 3} more mentions)"
            )
        else:
            return "\n".join(relevant_excerpts)

    async def analyze_location_connections(self) -> List[LocationConnection]:
        """Analyzes connections between locations in the project, processing in batches."""
        self.logger.info(
            f"Starting location connection analysis for project {self.project_id}"
        )
        try:
            locations_data = await db_instance.get_locations(
                self.user_id, self.project_id
            )
            if len(locations_data) < 2:
                self.logger.info("Not enough locations to analyze for connections.")
                return []

            self.logger.info(
                f"Found {len(locations_data)} locations to analyze for connections."
            )

            existing_connections_data = await db_instance.get_location_connections(
                self.project_id, self.user_id
            )
            existing_pairs = {
                tuple(sorted((c["location1_id"], c["location2_id"])))
                for c in existing_connections_data
            }

            all_new_connections = []
            batch_size = 25

            location_pairs = list(combinations(locations_data, 2))

            # Get all chapters for context retrieval
            chapters = await db_instance.get_all_chapters(self.user_id, self.project_id)
            chapter_content_map = {
                chapter["id"]: chapter["content"] for chapter in chapters
            }

            for i in range(0, len(location_pairs), batch_size):
                batch = location_pairs[i : i + batch_size]
                self.logger.info(
                    f"Processing batch {i//batch_size + 1} of location pairs."
                )

                batch_to_analyze = [
                    pair
                    for pair in batch
                    if tuple(sorted((pair[0]["id"], pair[1]["id"])))
                    not in existing_pairs
                ]

                if not batch_to_analyze:
                    continue

                # Enhanced formatting with RAG context
                formatted_pairs_with_context = []

                for idx, pair in enumerate(batch_to_analyze):
                    location1 = pair[0]
                    location2 = pair[1]

                    # Retrieve chapter context for each location
                    location1_context = await self._get_location_chapter_context(
                        location1["name"], chapter_content_map
                    )
                    location2_context = await self._get_location_chapter_context(
                        location2["name"], chapter_content_map
                    )

                    pair_info = f"Location Pair {idx+1}:\n"
                    pair_info += (
                        f"- Location 1: {location1['name']} (ID: {location1['id']})\n"
                    )
                    pair_info += f"  Description: {location1.get('description', 'No description')}\n"

                    if location1_context:
                        pair_info += f"  Context from chapters: {location1_context}\n"

                    pair_info += (
                        f"- Location 2: {location2['name']} (ID: {location2['id']})\n"
                    )
                    pair_info += f"  Description: {location2.get('description', 'No description')}\n"

                    if location2_context:
                        pair_info += f"  Context from chapters: {location2_context}\n"

                    formatted_pairs_with_context.append(pair_info)

                formatted_pairs = "\n\n".join(formatted_pairs_with_context)

                prompt = ChatPromptTemplate.from_template(
                    """You are a master world-builder. Your task is to identify meaningful connections between pairs of locations.
A connection might be geographical, political, cultural, historical, or based on travel routes.

Analyze the following pairs of locations and identify if a meaningful connection exists.
For each pair that is connected, explain the nature of this connection.
Use the provided context from chapters when available to help determine connections.

Location Pairs to Analyze:
{location_pairs}

Respond in JSON format with a single key "connections", which is a list of objects.
Each object must contain:
- "location1_id": The ID of the first location.
- "location2_id": The ID of the second location.
- "connection_type": A brief type for the connection (e.g., "Geographical", "Political", "Trade Route").
- "description": A detailed explanation of how the locations are connected.
- "travel_route": A description of the travel route, if applicable.
- "cultural_exchange": A description of cultural exchange, if applicable.

If a pair is not connected, do not include it in your response. If no connections are found, return an empty list.
"""
                )

                llm = await self._get_llm(self.model_settings["extractionLLM"])
                chain = prompt | llm | JsonOutputParser()

                try:
                    result = await chain.ainvoke({"location_pairs": formatted_pairs})
                    new_connections = result.get("connections", [])
                except Exception as e:
                    self.logger.error(
                        f"Error during location connection analysis LLM call: {e}"
                    )
                    continue

                for conn in new_connections:
                    location1_id = conn.get("location1_id")
                    location2_id = conn.get("location2_id")

                    if not location1_id or not location2_id:
                        continue

                    pair_key = tuple(sorted((location1_id, location2_id)))
                    if pair_key in existing_pairs:
                        continue

                    loc1_name = next(
                        (l["name"] for l in locations_data if l["id"] == location1_id),
                        "Unknown",
                    )
                    loc2_name = next(
                        (l["name"] for l in locations_data if l["id"] == location2_id),
                        "Unknown",
                    )

                    db_connection_id = await db_instance.create_location_connection(
                        location1_id=location1_id,
                        location2_id=location2_id,
                        location1_name=loc1_name,
                        location2_name=loc2_name,
                        connection_type=conn.get("connection_type", "Undefined"),
                        description=conn.get("description", ""),
                        travel_route=conn.get("travel_route"),
                        cultural_exchange=conn.get("cultural_exchange"),
                        project_id=self.project_id,
                        user_id=self.user_id,
                    )
                    self.logger.info(
                        f"Created connection between locations {location1_id} and {location2_id}."
                    )

                    # Add the name fields required by LocationConnection model
                    conn["id"] = db_connection_id
                    conn["location1_name"] = loc1_name
                    conn["location2_name"] = loc2_name

                    all_new_connections.append(LocationConnection(**conn))
                    existing_pairs.add(pair_key)

            self.logger.info(
                f"Location connection analysis complete. Found {len(all_new_connections)} new connections."
            )
            return all_new_connections
        except Exception as e:
            self.logger.error(
                f"An error occurred during location connection analysis: {e}",
                exc_info=True,
            )
            return []

    async def get_chat_history(self) -> List[Dict[str, Any]]:
        """Retrieves chat history for the current project."""
        self.logger.debug(f"Retrieving chat history for project {self.project_id}")
        try:
            # Directly call the database instance method
            history = await db_instance.get_chat_history(self.user_id, self.project_id)
            # Ensure it returns a list (db method should handle None case)
            return history if isinstance(history, list) else []
        except Exception as e:
            self.logger.error(f"Error retrieving chat history: {e}", exc_info=True)
            # Return empty list on error to avoid breaking chat interfaces
            return []

    async def reset_memory(self):
        """
        Resets chat memory by deleting all stored chat history from the database.
        This clears all conversation history for the current project.
        """
        try:
            # Delete chat history from database
            result = await db_instance.delete_chat_history(
                self.user_id, self.project_id
            )
            if result:
                self.logger.info(f"Chat history deleted for project {self.project_id}")
            else:
                self.logger.info(
                    f"No chat history found to delete for project {self.project_id}"
                )
            return result
        except Exception as e:
            self.logger.error(f"Error deleting chat history: {e}", exc_info=True)
            raise

    async def get_knowledge_base_content(self):
        """Gets all content from the vector store for this project."""
        if not self.vector_store:
            raise RuntimeError("Vector store not initialized.")
        return await self.vector_store.get_knowledge_base_content()

    # --- Text Extraction Utility ---
    async def extract_text_from_bytes(
        self, file_bytes: bytes, content_type: str
    ) -> Optional[str]:
        """Extract text content from file bytes."""
        try:
            if content_type.startswith("text/"):
                # Already text, just decode
                return file_bytes.decode("utf-8")

            # For PDFs
            if content_type == "application/pdf":
                # Parse PDF with PyPDF2
                from io import BytesIO
                from PyPDF2 import PdfReader

                pdf_file = BytesIO(file_bytes)
                reader = PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text

            # Add support for DOCX files
            if (
                content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                from io import BytesIO
                import docx

                docx_file = BytesIO(file_bytes)
                doc = docx.Document(docx_file)
                text = []
                for para in doc.paragraphs:
                    text.append(para.text)
                return "\n".join(text)

            # Add support for RTF files
            if content_type == "application/rtf" or content_type == "text/rtf":
                from striprtf.striprtf import rtf_to_text

                rtf_text = file_bytes.decode("utf-8", errors="ignore")
                return rtf_to_text(rtf_text)

            # Add support for plain text files with non-text MIME types
            if content_type == "application/octet-stream" and (
                file_bytes[:4].decode("ascii", errors="ignore").startswith("The")
                or all(
                    0x20 <= b <= 0x7E or b in (0x09, 0x0A, 0x0D)
                    for b in file_bytes[:100]
                )
            ):
                # Heuristic: if first 100 bytes look like ASCII text, treat as text
                return file_bytes.decode("utf-8", errors="replace")

            # Handle EPUB files
            if content_type == "application/epub+zip":
                from io import BytesIO
                import ebooklib
                from ebooklib import epub
                from bs4 import BeautifulSoup

                epub_file = BytesIO(file_bytes)
                book = epub.read_epub(epub_file)
                text = []

                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        # Extract HTML content and remove HTML tags
                        html_content = item.get_content().decode("utf-8")
                        soup = BeautifulSoup(html_content, "html.parser")
                        text.append(soup.get_text())

                return "\n\n".join(text)

            logging.warning(f"Unsupported content type: {content_type}")
            return None
        except Exception as e:
            logging.error(f"Error extracting text: {e}")
            return None

    # --- Utility Methods ---
    def chunk_content(self, content: str, max_chunk_size: int) -> List[str]:
        """Splits text into chunks based on estimated token count (simple version)."""
        if not content:
            return []
        # This is complex to do accurately without the tokenizer.
        # Simple splitting by paragraphs or sentences is often sufficient.
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        chunks = []
        current_chunk = ""
        for p in paragraphs:
            if not current_chunk:
                current_chunk = p
            elif (
                self.estimate_token_count(current_chunk + "\n\n" + p) <= max_chunk_size
            ):
                current_chunk += "\n\n" + p
            else:
                chunks.append(current_chunk)
                current_chunk = p
        if current_chunk:
            chunks.append(current_chunk)
        return chunks

    def chunk_content(self, content: str, max_chunk_size: int) -> List[str]:
        """Simple content chunker."""
        return [
            content[i : i + max_chunk_size]
            for i in range(0, len(content), max_chunk_size)
        ]

    async def get_proactive_suggestions(
        self, recent_chapters_content: str, notepad_content: str
    ) -> ProactiveSuggestionsResponse:
        """
        Generates proactive writing suggestions based on recent chapters and notepad content.
        """
        self.logger.info("Generating proactive writing suggestions.")
        try:
            parser = PydanticOutputParser(pydantic_object=ProactiveSuggestionsResponse)

            prompt = ChatPromptTemplate.from_template(
                """
                You are a proactive AI writing assistant. Your goal is to provide context-aware suggestions to enhance the user's writing.
                Analyze the provided recent chapter content and the user's notepad.
                Based on this context, provide a list of suggestions for enhancing sentences, paragraphs, or plot points.

                Recent Chapters Content:
                {recent_chapters_content}

                Notepad Content:
                {notepad_content}

                Respond in JSON format matching the following schema:
                {format_instructions}
                """
            )

            chain = prompt | self.check_llm | parser

            result = await chain.ainvoke(
                {
                    "recent_chapters_content": recent_chapters_content,
                    "notepad_content": notepad_content,
                    "format_instructions": parser.get_format_instructions(),
                }
            )

            return result

        except Exception as e:
            self.logger.error(
                f"Error generating proactive suggestions: {e}", exc_info=True
            )
            return ProactiveSuggestionsResponse(suggestions=[])

    async def extract_codex_items_from_text(self, content: str) -> List[Dict[str, Any]]:
        """
        A dedicated method to extract codex items from a given block of text.
        It encapsulates the logic of setting up a minimal state and running
        the extraction node, then saves the results to the database.
        """
        if not content.strip():
            self.logger.info("Content for codex extraction is empty. Skipping.")
            return []

        try:
            # 1. Get model settings
            model_settings = await self._get_model_settings()
            extraction_model_name = model_settings.get(
                "extractionLLM", "gemini-1.5-pro-002"
            )
            llm = await self._get_llm(extraction_model_name)

            # 2. Create a minimal state for the extraction node
            # Most fields are not required for this specific task.
            state = ChapterGenerationState(
                user_id=self.user_id,
                project_id=self.project_id,
                initial_chapter_content=content,
                llm=llm,
                # Set other required fields to default/None values
                chapter_number=0,
                plot="",
                writing_style="",
                instructions={},
                full_plot=None,
                plot_segment=None,
                total_chapters=0,
                check_llm=llm,
                vector_store=self.vector_store,
                summarize_chain=None,
            )

            # 3. Run the private extraction node
            extraction_result = await self._extract_codex_items_node(state)
            new_codex_items = extraction_result.get("new_codex_items", [])

            if not new_codex_items:
                self.logger.info(
                    "No new codex items were extracted from the provided text."
                )
                return []

            # 4. Save the extracted items to the database
            saved_items = []
            for item in new_codex_items:
                try:
                    # Use a single session for all creations within this call
                    async with db_instance.Session() as session:
                        item_id = await db_instance.create_codex_item(
                            name=item["name"],
                            description=item["description"],
                            type=item["type"],
                            subtype=item.get("subtype"),
                            user_id=self.user_id,
                            project_id=self.project_id,
                        )
                        saved_item_details = {
                            "id": item_id,
                            "name": item["name"],
                            "description": item["description"],
                            "type": item["type"],
                            "subtype": item.get("subtype"),
                        }
                        saved_items.append(saved_item_details)
                except Exception as db_error:
                    self.logger.error(
                        f"Error saving codex item '{item.get('name')}' to DB: {db_error}",
                        exc_info=True,
                    )
                    # Decide if we should continue or fail the whole batch
                    continue

            self.logger.info(
                f"Successfully extracted and saved {len(saved_items)} new codex items."
            )
            return saved_items

        except Exception as e:
            self.logger.error(
                f"An unexpected error occurred in extract_codex_items_from_text for project {self.project_id}: {e}",
                exc_info=True,
            )
            # Re-raise or handle as appropriate for the calling endpoint
            raise

    async def _get_location_chapter_context(
        self, location_name: str, chapter_content_map: Dict[str, str]
    ) -> str:
        """
        Retrieves relevant context from chapters that mention this location.
        Returns a concise summary of the mentions.
        """
        relevant_excerpts = []

        # Look for location name mentions in all chapters
        for chapter_id, content in chapter_content_map.items():
            if location_name in content:
                # Find the paragraph containing the location mention
                paragraphs = content.split("\n\n")
                for paragraph in paragraphs:
                    if location_name in paragraph:
                        # Truncate long paragraphs
                        if len(paragraph) > 300:
                            start_pos = max(0, paragraph.find(location_name) - 100)
                            excerpt = paragraph[start_pos : start_pos + 300] + "..."
                        else:
                            excerpt = paragraph

                        relevant_excerpts.append(excerpt)
                        break  # Just get one mention per chapter

        # Limit the total context length
        if not relevant_excerpts:
            return ""

        # Join up to 3 excerpts
        if len(relevant_excerpts) > 3:
            return (
                "\n".join(relevant_excerpts[:3])
                + f" (and {len(relevant_excerpts) - 3} more mentions)"
            )
        else:
            return "\n".join(relevant_excerpts)
